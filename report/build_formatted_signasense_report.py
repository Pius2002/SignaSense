from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from expand_signasense_report import arrow, box, canvas, font, generate_figures  # noqa: E402


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DOCX = BASE_DIR / "report.docx"
OUT_DOCX = BASE_DIR / "SignaSense_Report_Formatted_With_Diagrams.docx"
FIG_DIR = BASE_DIR / "generated_figures"

GROUP_MEMBERS = [
    ("NALUBOWA BENITAH MARGRET", "2023/BIT/169/PS"),
    ("ADEKE MARY", "2023/BIT/034/PS"),
    ("MUHUMUZA MUHAMMAD", "2023/BIT/252/PS"),
    ("MIREMBE JERMIMAH ARNN", "2023/BIT/148/PS"),
    ("KANSIIME PIUS", "2023/BIT/120/PS"),
]

SUPERVISOR = "Mr. Rogers Mwavu"

FIGURE_LIST = [
    "Figure 1: SignaSense conceptual framework.",
    "Figure 2: Use-case diagram for SignaSense users and modules.",
    "Figure 3: Android application navigation and accessibility flow.",
    "Figure 4: Prototype verification and testing workflow.",
    "Figure 5: Calibration and acceptance-testing workflow.",
    "Figure 6: SignaSense integrated system architecture.",
    "Figure 7: Smart glove flex-sensor voltage-divider wiring.",
    "Figure 8: Final smart glove prototype photograph placeholder.",
    "Figure 9: Smart stick ultrasonic sensor and buzzer wiring.",
    "Figure 10: Final smart stick prototype photograph placeholder.",
    "Figure 11: Camera sign backup recognition pipeline.",
    "Figure 12: User-defined sign training workflow.",
    "Figure 13: BLE data exchange between the smart glove and Android app.",
    "Figure 14: SignaSense deployment architecture.",
    "Figure 15: Smart stick left/right scanning procedure for route guidance.",
    "Figure 16: Smart stick distance filtering and guidance logic.",
    "Figure 17: SignaSense Android dashboard and accessibility controls.",
    "Figure 18: SignaSense project work plan Gantt chart.",
]

TABLE_LIST = [
    "Table 1: List of acronyms.",
    "Table 2: Summary of reviewed related studies.",
    "Table 3: Study population and planned sample size.",
    "Table 4: Data collection instruments.",
    "Table 5: User requirements.",
    "Table 6: Functional requirements.",
    "Table 7: System requirements.",
    "Table 8: Smart glove hardware pin allocation.",
    "Table 9: Smart stick hardware pin allocation.",
    "Table 10: Software tools and libraries used.",
    "Table 11: Functional test summary.",
    "Table 12: Project limitations and mitigation strategies.",
    "Table 13: Risk register.",
    "Table 14: Calibration record sheet.",
    "Table 15: Research work plan.",
    "Table 16: Literature gap analysis.",
    "Table 17: Prototype development phases.",
    "Table 18: Flex sensor calibration procedure.",
    "Table 19: Smart stick distance calibration procedure.",
    "Table 20: Glove BLE payload fields.",
    "Table 21: Camera backup operating conditions.",
    "Table 22: Test case matrix.",
    "Table 23: User acceptance checklist.",
]


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_document_defaults(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, "PAGE")
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def paragraph(doc, text="", align=None, bold=False, italic=False, size=12, style=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def bullet(doc, text):
    return paragraph(doc, f"- {text}")


def numbered(doc, text):
    p = paragraph(doc, f"{numbered.counter}. {text}")
    numbered.counter += 1
    return p


numbered.counter = 1


def shade(cell, color="EAF2F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ["top", "start", "bottom", "end"]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def table(doc, caption, headers, rows, widths=None, font_size=9.5):
    caption_p = paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    caption_p.paragraph_format.keep_with_next = True
    tbl = doc.add_table(rows=1, cols=len(headers))
    if "Table Grid" in [style.name for style in doc.styles if style.type == 3]:
        tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=font_size, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cells[i].paragraphs[0].add_run(str(val))
            set_run_font(run, size=font_size)
    if widths:
        for row in tbl.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width
    paragraph(doc, "")
    return tbl


def add_figure(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    return cap


def generate_extra_figures():
    figures = {}
    img, draw = canvas("SignaSense Use-Case Diagram")
    box(draw, (720, 180, 1110, 340), "SignaSense System", fill="#D9EAF7", fnt=font(28, bold=True))
    box(draw, (90, 150, 410, 285), "Blind User", fill="#DFF3E3", fnt=font(25, bold=True))
    box(draw, (90, 420, 410, 555), "Deaf User", fill="#DFF3E3", fnt=font(25, bold=True))
    box(draw, (90, 690, 410, 825), "Caregiver / Presenter", fill="#DFF3E3", fnt=font(23, bold=True))
    use_cases = [
        ((600, 470, 930, 585), "Receive obstacle\nguidance"),
        ((1010, 470, 1340, 585), "Read signs and\nbuild words"),
        ((600, 670, 930, 785), "Use camera\nbackup"),
        ((1010, 670, 1340, 785), "Train user-defined\nsigns"),
        ((1220, 180, 1600, 340), "Choose language\nand interface"),
    ]
    for xy, text in use_cases:
        box(draw, xy, text, fill="#FFF2CC", fnt=font(23, bold=True), radius=60)
    for start, end in [
        ((410, 215), (600, 525)),
        ((410, 485), (1010, 525)),
        ((410, 485), (600, 725)),
        ((410, 755), (1010, 725)),
        ((1110, 260), (1220, 260)),
    ]:
        arrow(draw, start, end)
    path = FIG_DIR / "figure_02_use_case.png"
    img.save(path)
    figures[2] = path

    img, draw = canvas("Smart Glove BLE Data Exchange")
    box(draw, (80, 210, 390, 360), "Flex sensors\nThumb to pinky", fill="#DFF3E3", fnt=font(25, bold=True))
    box(draw, (520, 210, 830, 360), "ESP32 glove\nADC smoothing\nletter state", fill="#D9EAF7", fnt=font(24, bold=True))
    box(draw, (960, 210, 1280, 360), "BLE status\nJSON payload", fill="#FFF2CC", fnt=font(24, bold=True))
    box(draw, (1390, 210, 1700, 360), "Android app\nword builder\nspeech output", fill="#F4DDE8", fnt=font(23, bold=True))
    box(draw, (520, 560, 830, 710), "Commands\ncalibrate\nclear\ncommit", fill="#EFEFEF", fnt=font(24, bold=True))
    box(draw, (960, 560, 1280, 710), "User actions\nhold sign\naccept word", fill="#EFEFEF", fnt=font(24, bold=True))
    for start, end in [((390, 285), (520, 285)), ((830, 285), (960, 285)), ((1280, 285), (1390, 285)), ((960, 635), (830, 635))]:
        arrow(draw, start, end)
    path = FIG_DIR / "figure_14_ble_data_exchange.png"
    img.save(path)
    figures[14] = path

    img, draw = canvas("SignaSense Deployment Architecture")
    box(draw, (100, 190, 470, 340), "Smart glove\nESP32 + sensors\nBLE peripheral", fill="#DFF3E3", fnt=font(23, bold=True))
    box(draw, (620, 190, 990, 340), "Android phone\nSignaSense app\nBLE central", fill="#D9EAF7", fnt=font(23, bold=True))
    box(draw, (1140, 190, 1510, 340), "User interface\ncaptions\nspeech\ntraining", fill="#FFF2CC", fnt=font(23, bold=True))
    box(draw, (100, 560, 470, 710), "Smart stick\nESP32 + ultrasonic\nlocal guidance", fill="#DFF3E3", fnt=font(23, bold=True))
    box(draw, (620, 560, 990, 710), "Bluetooth speaker\nSINOBAND Book\nvoice output", fill="#F4DDE8", fnt=font(23, bold=True))
    box(draw, (1140, 560, 1510, 710), "Blind user\nobstacle alerts\nmovement guide", fill="#FFF2CC", fnt=font(23, bold=True))
    for start, end in [((470, 265), (620, 265)), ((990, 265), (1140, 265)), ((470, 635), (620, 635)), ((990, 635), (1140, 635))]:
        arrow(draw, start, end)
    path = FIG_DIR / "figure_15_deployment_architecture.png"
    img.save(path)
    figures[15] = path

    img, draw = canvas("Calibration and Acceptance Testing Workflow")
    steps = [
        ((90, 190, 390, 325), "Power check\n3.3 V and GND"),
        ((500, 190, 800, 325), "Raw readings\nsensor by sensor"),
        ((910, 190, 1210, 325), "Open-hand\nbaseline"),
        ((1320, 190, 1620, 325), "Bent-hand\nbaseline"),
        ((90, 570, 390, 705), "Hold sign\nstable time"),
        ((500, 570, 800, 705), "Detect letter\nor reject noise"),
        ((910, 570, 1210, 705), "Build word\nconfirm"),
        ((1320, 570, 1620, 705), "User accepts\nspeech output"),
    ]
    for xy, text in steps:
        box(draw, xy, text, fill="#D9EAF7" if "Detect" not in text else "#FFF2CC", fnt=font(23, bold=True))
    for start, end in [((390, 257), (500, 257)), ((800, 257), (910, 257)), ((1210, 257), (1320, 257)), ((1620, 325), (1620, 500)), ((1620, 638), (1210, 638)), ((910, 638), (800, 638)), ((500, 638), (390, 638))]:
        arrow(draw, start, end)
    path = FIG_DIR / "figure_16_calibration_acceptance.png"
    img.save(path)
    figures[16] = path
    return figures


def add_photo_placeholder(doc, caption, note):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if "Table Grid" in [style.name for style in doc.styles if style.type == 3]:
        tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade(cell, "FFF2CC")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(note)
    set_run_font(r, size=11, bold=True)
    paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)


def add_toc_field(doc):
    toc_lines = [
        ("ACKNOWLEDGEMENT", "4"),
        ("INTRODUCTION", "11"),
        ("CONCEPT PAPER", "12"),
        ("Concept Paper Format", "12"),
        ("1) Face page", "12"),
        ("2) Introduction", "12"),
        ("3) Background", "12"),
        ("4) Problem Statement", "12"),
        ("5) Main Objective/aim", "12"),
        ("6) Specific Objectives", "13"),
        ("7) Scope", "13"),
        ("8) Methodology", "13"),
        ("PROJECT PROPOSAL", "15"),
        ("CHAPTER ONE: INTRODUCTION", "15"),
        ("1.1 Introduction", "15"),
        ("1.2 Background", "15"),
        ("1.3 Problem statement", "15"),
        ("1.4 General Objective /Aim/ Purpose", "15"),
        ("1.5 Specific Objectives", "16"),
        ("1.6 Research questions", "16"),
        ("1.7 Significance / Justification (choose one)", "16"),
        ("1.8 Scope", "16"),
        ("CHAPTER TWO: LITERATURE REVIEW", "18"),
        ("What is a literature review?", "18"),
        ("2.1 Introduction", "18"),
        ("CHAPTER 3: METHODOLOGY", "22"),
        ("3.0: Introduction", "22"),
        ("3.1 Sampling techniques", "22"),
        ("3.2 Study population and sample size", "22"),
        ("3.3 Data Collection Procedure", "22"),
        ("3.4: Data Collection Instruments", "23"),
        ("3.5 Data processing and analysis", "23"),
        ("3.6 Data presentation and interpretation", "23"),
        ("3.7 Limitation of the study", "23"),
        ("REFERENCES", "27"),
        ("APPENDICES", "27"),
        ("THE REPORT", "28"),
        ("CHAPTER FOUR: DATA PRESENTATION, ANALYSIS AND INTERPRETATION", "28"),
        ("4.1 Introduction", "28"),
        ("CHAPTER FIVE: SYSTEM DEVELOPMENT AND IMPLEMENTATION", "32"),
        ("5.1. Requirements identified", "32"),
        ("5.1.1 User Requirements", "32"),
        ("5.1.2 Functional Requirements", "32"),
        ("5.2 System requirements (of the new system)", "33"),
        ("5.3 Architectural design of the new system", "33"),
        ("CHAPTER SIX: SUMMARY, CONCLUSION AND RECOMMENDATIONS", "44"),
        ("6.1 Summary (optional)", "44"),
        ("6.2 Conclusion", "44"),
        ("6.3 Recommendations", "45"),
        ("6.4 Areas for further study", "45"),
        ("REFERENCES", "46"),
        ("APPENDICES", "47"),
        ("INDICES", "47"),
        ("APPENDIX A: ASSESMENT FORM", "48"),
        ("APPENDIX B: TYPOGRAPHIC FORMAT", "50"),
        ("APPENDIX C: SUPERVISION MONITORING FORM", "50"),
        ("APPENDIX D: RESEACH WORKPLAN", "52"),
        ("APPENDIX E: FACE PAGE", "54"),
    ]
    for title, page in toc_lines:
        p = paragraph(doc, "")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, size=11)


def add_title_page(doc):
    paragraph(doc, "MBARARA UNIVERSITY OF SCIENCE AND TECHNOLOGY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    paragraph(doc, "INSTITUTE OF COMPUTER SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "DEPARTMENT OF INFORMATION TECHNOLOGY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    paragraph(doc, "")
    paragraph(doc, "SIGNASENSE - AN INTEGRATED ASSISTIVE SOLUTION FOR THE DEAF AND BLIND", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    paragraph(doc, "")
    paragraph(doc, "BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for name, reg in GROUP_MEMBERS:
        paragraph(doc, f"{name}    {reg}", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "")
    paragraph(doc, "A final year project report submitted to the Institute of Computer Science in partial fulfilment of the requirements for the award of the Bachelor of Science Degree in Information Technology of Mbarara University of Science and Technology.", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "")
    paragraph(doc, f"Supervisor: {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc, "Department of Computer Science, Institute of Computer Science, Mbarara University of Science and Technology", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "May, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_front_matter(doc):
    heading(doc, "Declaration", 1)
    paragraph(doc, "We declare that this report is our original work and has not been submitted to any institution for any academic award. Where other people's work has been used, it has been properly acknowledged through in-text citations and references.")
    table(doc, "Student declaration signatures", ["Name", "Registration Number", "Signature", "Date"], [(n, r, "________________________", "________________") for n, r in GROUP_MEMBERS], font_size=9)
    doc.add_page_break()
    heading(doc, "Approval", 1)
    paragraph(doc, "This report has been submitted for examination with the approval of the project supervisor.")
    paragraph(doc, f"Supervisor: {SUPERVISOR}")
    paragraph(doc, "Signature: ________________________________        Date: __________________")
    doc.add_page_break()
    heading(doc, "ACKNOWLEDGEMENT", 1)
    paragraph(doc, f"We express our sincere gratitude to our supervisor, {SUPERVISOR}, for his guidance, support and constructive feedback throughout the development of this project. We also appreciate the Department of Information Technology and the Institute of Computer Science for providing the academic structure and project supervision process that guided this work.")
    paragraph(doc, "We thank our colleagues, classmates and all persons who supported the testing and discussion of SignaSense during development. Their observations helped improve the smart glove, smart stick, Android application, Bluetooth connection workflow and camera sign backup.")
    paragraph(doc, "We finally thank our families and friends for their encouragement throughout the project period.")
    doc.add_page_break()
    heading(doc, "Abstract", 1)
    paragraph(doc, "People with hearing and visual impairments face communication, navigation and social participation barriers. SignaSense was developed as a low-cost assistive prototype that combines a smart glove for sign-to-text/speech support, a smart stick for obstacle detection and spoken guidance, and an Android application for onboarding, captions, speech output, camera sign backup and user-defined sign training.")
    paragraph(doc, "The smart glove uses ESP32 firmware to read flex sensors, calibrate finger bends and send glove status to the phone over Bluetooth Low Energy. The smart stick uses an ultrasonic sensor to detect obstacles, speaks distance and guidance through a Bluetooth speaker, and keeps a buzzer as backup. The Android app supports blind and deaf onboarding, English/Luganda/Acholi language selection, dark/light mode, smart glove BLE, word and sentence suggestions, camera sign recognition and private user-defined signs.")
    paragraph(doc, "The project followed an iterative prototype methodology involving requirements analysis, hardware design, Arduino IDE firmware development, Android development, calibration, serial testing, phone installation and functional verification. Testing confirmed that the app and camera model load correctly, BLE communication works, and the stick guidance logic can be configured for distance filtering and route advice. Hardware limitations such as flex-sensor voltage-divider stability and ESP32 ADC pin behaviour were documented for future improvement.")
    doc.add_page_break()
    heading(doc, "TABLE OF CONTENTS", 1)
    add_toc_field(doc)
    doc.add_page_break()
    heading(doc, "List of Figures", 1)
    for item in FIGURE_LIST:
        paragraph(doc, item)
    doc.add_page_break()
    heading(doc, "List of Tables", 1)
    for item in TABLE_LIST:
        paragraph(doc, item)
    doc.add_page_break()
    heading(doc, "List of Acronyms", 1)
    table(doc, "Table 1: List of acronyms.", ["Acronym", "Meaning"], [
        ("ADC", "Analogue-to-Digital Converter"),
        ("A2DP", "Advanced Audio Distribution Profile"),
        ("API", "Application Programming Interface"),
        ("BLE", "Bluetooth Low Energy"),
        ("ESP32", "Espressif 32-bit Wi-Fi/Bluetooth microcontroller"),
        ("GPIO", "General Purpose Input/Output"),
        ("ICT", "Information and Communication Technology"),
        ("TTS", "Text-to-Speech"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
    ])
    doc.add_page_break()


def add_introduction(doc):
    heading(doc, "INTRODUCTION", 1)
    paragraph(doc, "This document presents the full SignaSense project in the structure used by the Institute of Computer Science undergraduate research format. It begins with the concept paper, continues with the project proposal chapters and then presents the final report chapters on implementation, results, conclusion and recommendations.")
    paragraph(doc, "The report documents the system built during the project: an ESP32 smart glove, an ESP32 smart stick and the SignaSense Android application. The main design concern was practical assistive support for deaf and blind users using affordable hardware, offline operation where possible and simple interaction flows.")
    paragraph(doc, "The report also includes diagrams, architectures, wiring designs, workflows, test summaries, risk management and appendices required for submission.")


def add_concept_paper(doc, figures):
    doc.add_page_break()
    heading(doc, "CONCEPT PAPER", 1)
    heading(doc, "Concept Paper Format", 2)
    paragraph(doc, "The concept paper summarises the project idea before full development. It identifies the problem, background, objectives, scope and methodology that guided the SignaSense prototype.")
    heading(doc, "1) Face page", 3)
    paragraph(doc, "Project title: SignaSense - An Integrated Assistive Solution for the Deaf and Blind.")
    paragraph(doc, "Student group: NALUBOWA BENITAH MARGRET, ADEKE MARY, MUHUMUZA MUHAMMAD, MIREMBE JERMIMAH ARNN and KANSIIME PIUS.")
    paragraph(doc, f"Supervisor: {SUPERVISOR}.")
    heading(doc, "2) Introduction", 3)
    paragraph(doc, "SignaSense is an assistive technology prototype designed to support communication for deaf users and navigation for blind users. The project combines a glove, a walking stick and a mobile application so that one system can serve both communication and mobility needs.")
    heading(doc, "3) Background", 3)
    paragraph(doc, "Assistive technologies such as smart gloves, camera sign-recognition systems, walking sticks and mobile accessibility applications are increasingly used to support persons with disabilities. However, many products are expensive, require internet services or focus on only one disability. In Uganda, disability inclusion remains important because accessible communication and assistive devices are recognised in law and policy (Uganda, 2020).")
    paragraph(doc, "The project therefore explores a locally maintainable design using ESP32 microcontrollers, flex sensors, an ultrasonic sensor, Bluetooth communication and an Android application.")
    heading(doc, "4) Problem Statement", 3)
    paragraph(doc, "People with hearing and visual impairments often lack affordable integrated tools that support both communication and safe mobility. A deaf user may need caption and speech support, while a blind user may need obstacle warnings and spoken guidance. Existing single-function devices do not fully address the combined problem.")
    heading(doc, "5) Main Objective/aim", 3)
    paragraph(doc, "To develop a low-cost integrated assistive prototype that improves communication support for deaf users and mobility guidance for blind users.")
    heading(doc, "6) Specific Objectives", 3)
    numbered(doc, "To examine existing assistive systems for sign communication and obstacle detection.")
    numbered(doc, "To design and implement a smart glove that detects finger bend patterns and sends recognised letters or words to a mobile phone.")
    numbered(doc, "To design and implement a smart stick that detects obstacles, speaks distance and gives path guidance.")
    numbered(doc, "To develop an Android application that supports onboarding, captions, speech, camera sign backup, word suggestions and user-defined signs.")
    numbered(doc, "To test the prototype modules and document limitations for future improvement.")
    heading(doc, "7) Scope", 3)
    paragraph(doc, "The project scope covers prototype design, wiring, firmware, Android app development and functional testing. It focuses on one-hand alphabet and useful word formation for the glove, obstacle-distance and left/right guidance for the stick, and phone-based visual/audio interfaces for users. Full natural sign-language grammar, certified medical-device testing and large-scale field deployment are outside the scope.")
    heading(doc, "8) Methodology", 3)
    paragraph(doc, "The project used an iterative prototyping methodology. The team identified user requirements, selected affordable components, assembled circuits, wrote Arduino IDE firmware, built the Android app, tested communication and refined behaviour based on observed faults such as random letters, zero flex readings, Bluetooth delay and ultrasonic reading instability.")
    add_figure(doc, FIG_DIR / "figure_01_conceptual_framework.png", "Figure 1: SignaSense conceptual framework.")


def add_project_proposal(doc, figures):
    doc.add_page_break()
    heading(doc, "PROJECT PROPOSAL", 1)
    heading(doc, "CHAPTER ONE: INTRODUCTION", 1)
    heading(doc, "1.1 Introduction", 2)
    paragraph(doc, "People with sensory disabilities face challenges that affect communication, safe movement and social participation. Hearing loss can limit speech communication and education access, while visual impairment can reduce independent mobility and access to daily activities. SignaSense responds by integrating glove-based communication, camera sign backup and smart-stick guidance in one project.")
    paragraph(doc, "The World Health Organization reports that large numbers of people live with hearing loss and vision impairment globally, creating a continuing need for accessible and affordable assistive technologies (WHO, 2026a; WHO, 2026b).")
    heading(doc, "1.2 Background", 2)
    paragraph(doc, "Assistive technology has advanced through wearable sensors, embedded microcontrollers, mobile computing and computer vision. ESP32 boards are suitable for low-cost prototyping because they support Bluetooth and analogue/digital interfaces, while Android phones provide TextToSpeech, SpeechRecognizer, CameraX and local storage APIs (Android Developers, 2026a; Android Developers, 2026b; Android Developers, 2026c).")
    paragraph(doc, "The project moved the glove from Wi-Fi to BLE because ESP32 ADC2 pins can conflict with Wi-Fi, while BLE provided faster app communication and reduced delay. The stick was returned to Bluetooth speaker output because real-time mobility guidance needs immediate speech without phone polling delay.")
    heading(doc, "1.3 Problem statement", 2)
    paragraph(doc, "People with hearing and visual impairments in low-resource settings have limited access to affordable, locally maintainable assistive systems that combine communication and mobility support. Existing tools are often expensive, require stable internet access, or solve only one part of the problem.")
    heading(doc, "1.4 General Objective /Aim/ Purpose", 2)
    paragraph(doc, "To enhance communication and mobility support for deaf and blind users through an integrated low-cost SignaSense prototype.")
    heading(doc, "1.5 Specific Objectives", 2)
    numbered(doc, "To review existing technologies for sign-language translation, obstacle detection and mobile accessibility.")
    numbered(doc, "To design and build a smart glove that reads flex-sensor bend patterns and communicates with a phone.")
    numbered(doc, "To design and build a smart stick that measures obstacle distance and gives spoken guidance.")
    numbered(doc, "To develop an Android app that supports blind/deaf onboarding, captions, speech, BLE glove data, camera sign backup and user-defined signs.")
    numbered(doc, "To test the system and document results, limitations and future improvements.")
    heading(doc, "1.6 Research questions", 2)
    numbered(doc, "What gaps exist in current communication and mobility assistive technologies?")
    numbered(doc, "How can affordable components be integrated into one assistive prototype?")
    numbered(doc, "How reliable are the glove, stick and app under prototype testing?")
    numbered(doc, "How can mobile app features improve accessibility and reduce communication errors?")
    heading(doc, "1.7 Significance / Justification (choose one)", 2)
    paragraph(doc, "The project is significant because it combines two assistive functions that are often treated separately. The glove supports captions and speech from signs, while the stick supports obstacle awareness and navigation. The camera backup and user-defined signs improve resilience when the glove battery is low or a user cannot make standard signs.")
    paragraph(doc, "The system is also educationally significant because it demonstrates practical embedded-system design, Android development, Bluetooth communication, sensor calibration and human-centred accessibility design.")
    heading(doc, "1.8 Scope", 2)
    paragraph(doc, "The study covers the design and testing of a prototype made from an ESP32 smart glove, ESP32 smart stick and Android application. It includes hardware wiring, software development, Bluetooth communication, camera model integration, app UI/UX and functional testing. The study does not cover mass production, clinical certification or full continuous sign-language grammar.")
    if (FIG_DIR / "figure_02_use_case.png").exists():
        add_figure(doc, FIG_DIR / "figure_02_use_case.png", "Figure 2: Use-case diagram for SignaSense users and modules.")
    add_figure(doc, FIG_DIR / "figure_08_app_navigation.png", "Figure 3: Android application navigation and accessibility flow.")


def add_literature_review(doc):
    doc.add_page_break()
    heading(doc, "CHAPTER TWO: LITERATURE REVIEW", 1)
    heading(doc, "What is a literature review?", 2)
    paragraph(doc, "A literature review studies existing work related to the project problem, identifies strengths and gaps, and shows how the new system contributes to the field. For SignaSense, the literature review covers smart gloves, smart walking sticks, mobile accessibility, camera sign recognition and inclusive technology.")
    heading(doc, "2.1 Introduction", 2)
    paragraph(doc, "Assistive technologies for hearing and visual impairments include wearable sign-recognition gloves, camera-based hand recognition, smart walking sticks and mobile accessibility applications. These systems can improve independence, but many existing prototypes work in isolation.")
    heading(doc, "2.2 Review of literature related to objective 1", 2)
    paragraph(doc, "Wearable sensor gloves use flex sensors, inertial sensors or contact sensors to capture hand movement. Their advantage is that they are less affected by background and lighting than camera-only systems. However, flex sensors alone cannot distinguish every alphabet sign because some signs depend on palm orientation, finger crossing and motion.")
    heading(doc, "2.3 Review of literature related to objective 2", 2)
    paragraph(doc, "Smart walking sticks commonly use ultrasonic or infrared sensors to detect obstacles. The HC-SR04 type ultrasonic sensor measures distance by sending a sound pulse and timing the returning echo. SignaSense extends basic distance detection by adding speech, calibrated distance output, stop warnings and left/right scan guidance.")
    heading(doc, "2.4 Review of literature related to objective 3", 2)
    paragraph(doc, "Mobile phones support accessible output through text display, vibration, speech and camera processing. Android TextToSpeech enables spoken output and SpeechRecognizer can support voice commands where device services are available. CameraX allows image-analysis pipelines, and MediaPipe Hand Landmarker supports hand landmark detection on Android (Google AI Edge, 2025).")
    heading(doc, "2.5 Any other relevant literature", 2)
    paragraph(doc, "Uganda's Persons with Disabilities Act recognises assistive devices and accessible communication, including sign language, audio, visual and accessible ICT formats (Uganda, 2020). This supports the relevance of a locally developed assistive prototype.")
    heading(doc, "2.6 General remarks", 2)
    paragraph(doc, "The reviewed work shows that no single low-cost prototype fully solves sign communication, word output, camera backup and obstacle guidance together. SignaSense therefore contributes an integrated, modular approach.")
    table(doc, "Table 2: Summary of reviewed related studies.", ["Area", "Strength", "Gap", "SignaSense response"], [
        ("Smart glove", "Captures finger movement directly", "Limited full-language recognition", "Uses deliberate alphabet, word building and app suggestions."),
        ("Camera sign recognition", "Works without glove hardware", "Sensitive to light and camera position", "Added as backup, not replacement."),
        ("Smart walking stick", "Detects nearby obstacles", "Often only beeps or gives basic warnings", "Speaks distance and guides left/right scanning."),
        ("Mobile accessibility app", "Can provide captions and speech", "May depend on internet or external devices", "Keeps core features local on phone."),
    ], font_size=9)
    heading(doc, "2.7 Literature gap analysis", 2)
    paragraph(doc, "The most important gap identified from the reviewed work is integration. A deaf user mainly needs expressive communication support, while a blind user mainly needs mobility and environmental awareness. In practice, disability support devices are often designed as separate products, which increases cost and makes demonstration, maintenance and user training harder. SignaSense addresses this by combining a smart glove, a smart stick and a common Android interface under one project.")
    paragraph(doc, "A second gap is feedback control. Many student-level smart stick prototypes use a buzzer only. A buzzer can warn that something is close, but it does not explain distance, direction or the next action. For a blind user, spoken guidance is more informative than a tone because it can distinguish between warning, danger and movement instruction. This is why the smart stick was designed to speak an obstacle warning, wait for the current phrase to finish, guide the user to scan left and right, then advise the side with more free space.")
    paragraph(doc, "A third gap is personalisation. Standard alphabets are useful, but not every deaf user can make every hand shape. Some users may have missing fingers, weak finger movement, injury or different signing habits. The user-defined sign section was therefore added so that a phone can locally learn a person's own signs without changing the general model for other users.")
    table(doc, "Table 16: Literature gap analysis.", ["Observed gap", "Effect on users", "Design response in SignaSense"], [
        ("Separate devices for separate disabilities", "Higher cost and fragmented use", "One project combines glove, stick and app."),
        ("Buzzer-only obstacle warning", "User knows danger exists but not what to do", "Speech gives distance and movement guidance."),
        ("Camera-only sign recognition limitations", "Lighting and background can reduce accuracy", "Camera mode is a backup while glove BLE remains primary."),
        ("Fixed sign alphabet", "Users with hand differences may be excluded", "Local user-defined signs can be trained on one phone."),
        ("Developer diagnostics visible to users", "Interfaces become confusing", "The final UI hides raw diagnostics from normal users."),
    ], font_size=9)
    heading(doc, "2.8 Theoretical basis for the project", 2)
    paragraph(doc, "The glove component is based on the principle of variable resistance. A flex sensor changes resistance when bent. When it is connected in a voltage divider, the ESP32 analogue input reads a changing voltage that can be mapped into a straight, half-bent or bent state. The accuracy of this method depends on a stable 3.3 V supply, common ground, suitable resistor values and calibration for each finger.")
    paragraph(doc, "The stick component is based on time-of-flight measurement. The ultrasonic sensor sends a sound pulse and waits for the echo. The firmware converts echo time into distance using the speed of sound, then applies calibration scale and offset. Invalid readings are rejected because ultrasonic sensors can return zero, timeout values or unstable values when the surface is soft, angled or outside the range.")
    paragraph(doc, "The app component is based on human-computer interaction principles for accessibility. A blind user should not depend on small visual controls, so audio prompts and speech output are needed. A deaf user needs readable captions, word suggestions and clear confirmation before speech. The dark and light mode settings support different environments and reduce eye strain.")
    heading(doc, "2.9 Summary of the literature review", 2)
    paragraph(doc, "The literature supports the technical direction of the project but also shows why the system must be presented honestly. Flex sensors can detect bends but cannot identify every international sign without extra orientation and motion data. Camera recognition can support backup detection but needs good lighting and a trained model. Bluetooth speakers can output audio, but they do not normally provide voice command input unless they expose a microphone profile that the system can use. These realities guided the final design choices.")


def add_methodology(doc):
    doc.add_page_break()
    heading(doc, "CHAPTER 3: METHODOLOGY", 1)
    heading(doc, "3.0: Introduction", 2)
    paragraph(doc, "This chapter describes the research and development methods used to build SignaSense. The project combined descriptive research, prototype development and functional testing.")
    heading(doc, "3.1 Sampling techniques", 2)
    paragraph(doc, "Purposive sampling was planned because the project targets users with specific communication or mobility needs. The intended respondents include deaf users, blind users, disability support staff, supervisors and student testers.")
    heading(doc, "3.2 Study population and sample size", 2)
    paragraph(doc, "The planned study population included persons with hearing and visual impairments around Mbarara Municipality and university stakeholders involved in accessibility support. Formal participant testing should be completed before field deployment.")
    table(doc, "Table 3: Study population and planned sample size.", ["Category", "Population", "Planned sample"], [
        ("Deaf users", "Users who communicate through signs", "5"),
        ("Blind users", "Users who need obstacle guidance", "5"),
        ("Disability support staff", "Teachers, support officers or caregivers", "3"),
        ("Student testers", "Prototype demonstration testers", "5"),
    ], font_size=9)
    heading(doc, "3.3 Data Collection Procedure", 2)
    paragraph(doc, "The procedure involved identifying the project problem, reviewing related work, designing circuits, writing firmware, developing the Android app, testing each module, recording observed faults and improving the system.")
    heading(doc, "3.4: Data Collection Instruments", 2)
    table(doc, "Table 4: Data collection instruments.", ["Instrument", "Purpose"], [
        ("Questionnaire", "Collect structured feedback on usefulness, comfort, accuracy and clarity."),
        ("Interview guide", "Collect deeper comments from users and support staff."),
        ("Observation checklist", "Record wiring faults, app delay, BLE connection and sensor behaviour."),
        ("Serial monitor logs", "Verify raw readings, bend states, distance and connection status."),
    ], font_size=9)
    heading(doc, "3.5 Data processing and analysis", 2)
    paragraph(doc, "Quantitative data would be processed using counts, percentages, average response time and distance error. Qualitative feedback would be grouped into themes such as ease of use, comfort, safety, audio clarity, sign accuracy and suggested improvements.")
    heading(doc, "3.6 Data presentation and interpretation", 2)
    paragraph(doc, "Data is presented using narrative explanations, tables, screenshots, wiring diagrams, system architecture diagrams and workflow diagrams. Test results are interpreted against the objectives and user requirements.")
    heading(doc, "3.7 Limitation of the study", 2)
    paragraph(doc, "The major limitations were limited formal participant testing, flex-sensor wiring instability, ADC pin restrictions on ESP32, ultrasonic sensor noise, Bluetooth speaker connection constraints and the fact that full sign-language grammar cannot be achieved using flex sensors alone.")
    heading(doc, "3.8 Prototype development method", 2)
    paragraph(doc, "The project followed an iterative prototype method. Each module was first built independently, then integrated after its basic behaviour was confirmed. The smart stick was tested from simple buzzer alerts to Bluetooth speaker voice output. The glove moved from raw analogue readings to BLE status messages and Android word building. The Android application was refined from a developer test screen into a cleaner user-facing interface.")
    table(doc, "Table 17: Prototype development phases.", ["Phase", "Main activity", "Output"], [
        ("Phase 1", "Problem analysis and concept paper preparation", "Approved SignaSense concept."),
        ("Phase 2", "Hardware pin mapping and basic Arduino sketches", "ESP32 glove and stick test firmware."),
        ("Phase 3", "Android app onboarding, glove view and stick view", "Installable SignaSense APK."),
        ("Phase 4", "BLE and Wi-Fi comparison", "BLE retained for glove; stick returned to local Bluetooth speaker mode."),
        ("Phase 5", "Camera backup and user-defined signs", "General and custom sign recognition options."),
        ("Phase 6", "Report writing, diagrams and verification", "Formatted final report and rendered pages."),
    ], font_size=9)
    heading(doc, "3.9 Hardware construction procedure", 2)
    paragraph(doc, "For the glove, five flex sensors are placed along the fingers. Each sensor must be connected as a voltage divider using 3.3 V, a resistor to ground and the centre point to the ESP32 analogue pin. All sensor grounds must join the ESP32 ground. The final pin status used for the glove is thumb GPIO25, index GPIO33, middle GPIO32, ring GPIO35 and pinky GPIO34.")
    paragraph(doc, "For the smart stick, the ultrasonic trigger pin is connected to GPIO5, echo is connected to GPIO18 through a level shifter or voltage divider when using a 5 V ultrasonic module, and the buzzer is connected to GPIO15 as backup. The Bluetooth speaker named SINOBAND Book is used for spoken output when connected.")
    heading(doc, "3.10 Calibration procedure", 2)
    paragraph(doc, "Calibration was treated as a required part of the method because both flex sensors and ultrasonic sensors vary between components. Flex sensors need open-hand and bent-hand baselines. Ultrasonic distance needs ruler or tape-measure checking at known distances and adjustment through scale and offset constants.")
    table(doc, "Table 18: Flex sensor calibration procedure.", ["Step", "Action", "Acceptance condition"], [
        ("1", "Power the ESP32 from a stable source", "Serial monitor shows non-zero raw readings."),
        ("2", "Keep the hand open and record all five readings", "Each finger has a repeatable open baseline."),
        ("3", "Bend one finger at a time and record readings", "The active finger changes more than noise."),
        ("4", "Save straight and bent values in firmware/app logic", "Bend percentage moves smoothly from straight to bent."),
        ("5", "Hold each target letter five times", "The same letter is detected when held steadily."),
    ], font_size=9)
    table(doc, "Table 19: Smart stick distance calibration procedure.", ["Step", "Action", "Acceptance condition"], [
        ("1", "Measure a flat obstacle at 30 cm", "Serial output is close to 30 cm after filtering."),
        ("2", "Measure at 50 cm and 100 cm", "Error remains consistent and can be corrected."),
        ("3", "Adjust temperature, scale and offset constants", "Displayed metres and centimetres match the tape measure."),
        ("4", "Test angled and soft objects", "Invalid readings are rejected rather than spoken."),
        ("5", "Test left and right scanning", "The stick recommends the side with greater measured clearance."),
    ], font_size=9)
    heading(doc, "3.11 Testing strategy", 2)
    paragraph(doc, "Testing was divided into unit testing, integration testing and user acceptance preparation. Unit testing checked one sensor or one software function at a time. Integration testing checked BLE messages, app display, camera mode and smart-stick audio output. Acceptance testing focuses on whether a user can understand the output and operate the prototype without developer help.")
    paragraph(doc, "The testing strategy also includes fault testing. If all glove readings are zero, the system should warn that no glove signal is present. If all values are too high, the wiring should be checked. If a letter appears for less than the stability time, it should be rejected as noise. If the stick reading is invalid or missing, speech should not be repeated blindly.")
    heading(doc, "3.12 Ethical considerations", 2)
    paragraph(doc, "Testing with deaf and blind users requires informed consent and respect. Participants must be told that SignaSense is a prototype, not a certified medical or safety device. Camera samples and user-defined signs should remain on the local phone unless the user agrees to share them. The smart stick should be tested in controlled areas before being used in open public spaces.")


def add_report_chapters(doc, figures):
    doc.add_page_break()
    heading(doc, "REFERENCES", 1)
    paragraph(doc, "Proposal references are consolidated in the final Harvard-style reference list after Chapter Six.")
    heading(doc, "APPENDICES", 1)
    paragraph(doc, "Proposal appendices are consolidated with the final report appendices.")
    doc.add_page_break()
    heading(doc, "THE REPORT", 1)
    heading(doc, "CHAPTER FOUR: DATA PRESENTATION, ANALYSIS AND INTERPRETATION", 1)
    heading(doc, "4.1 Introduction", 2)
    paragraph(doc, "This chapter presents and interprets development data collected from prototype construction, serial monitor testing, Android installation, BLE checks, camera model integration and smart-stick behaviour testing.")
    heading(doc, "4.1 Presentation and interpretation of results on objective (research question 1)", 2)
    paragraph(doc, "The review of existing systems showed that most assistive systems focus on one problem at a time. Smart gloves translate signs but may not support mobility. Smart sticks support navigation but do not support communication. The finding justified an integrated system.")
    heading(doc, "4.2 Presentation and interpretation of results on objective (research question 2)", 2)
    paragraph(doc, "The design objective was achieved through the smart glove, smart stick and Android application. The glove connects through BLE, the stick uses ultrasonic sensing and Bluetooth speaker output, and the app provides visual and audio interaction modes.")
    heading(doc, "4.3 Presentation and interpretation of results on objective (research question 3)", 2)
    paragraph(doc, "Testing showed that the Android application builds and installs successfully, the camera model loads, BLE communication can be established, and the stick logic can speak only when obstacles require guidance. The main hardware issue observed during glove testing was zero or unstable flex-sensor readings when the voltage divider was not producing a valid analogue signal.")
    table(doc, "Table 11: Functional test summary.", ["Module", "Test", "Expected result", "Observed result"], [
        ("Android app", "Install APK on phone", "App opens as SignaSense", "Successful."),
        ("Camera mode", "Load hand model", "Camera backup opens", "Model assets load; cleaned UI hides diagnostics."),
        ("Glove BLE", "Scan and connect", "Phone receives glove status", "BLE status data received."),
        ("Glove sensors", "Read five bends", "Non-zero values change with bends", "Requires stable voltage-divider wiring; zero readings documented."),
        ("Smart stick", "Obstacle scan", "Speak warning and guidance", "Firmware designed with distance filtering and left/right scan procedure."),
        ("Bluetooth speaker", "Connect SINOBAND Book", "Voice output through speaker", "Configured by name and Bluetooth output mode."),
    ], font_size=9)
    add_figure(doc, FIG_DIR / "figure_12_testing_workflow.png", "Figure 4: Prototype verification and testing workflow.")
    heading(doc, "4.4 Analysis of smart glove sensor results", 2)
    paragraph(doc, "The glove testing process showed that bend detection depends more on circuit quality than on the software alone. When a flex sensor is connected without a resistor, the ESP32 input can float or remain near zero, which makes reliable bend detection impossible. When a proper voltage divider is used, bending changes the analogue voltage and the app can classify finger states.")
    paragraph(doc, "The observed unstable readings were useful because they exposed the difference between raw electrical input and meaningful sign detection. The final design therefore treats raw readings as diagnostic data inside the system, but the user interface displays only useful outputs such as current letter, word suggestions and speech actions.")
    paragraph(doc, "The glove app was also changed to avoid immediately committing random letters. A letter must remain stable for the configured hold time before it is allowed into the word area. This improves correctness because short spikes are treated as noise rather than communication.")
    heading(doc, "4.5 Analysis of smart stick results", 2)
    paragraph(doc, "The smart stick testing showed that speaking every distance value causes repeated audio and confuses the user. The firmware was therefore refined so that it speaks only when an obstacle is inside the warning range, gives a short obstacle warning, then gives the action. For very close obstacles, it says stop now and starts the left/right scan procedure.")
    paragraph(doc, "The left/right procedure gives the user time to move the stick. The system requests a left scan, waits while samples are collected, requests a right scan, waits again, compares the best distances and recommends the safer side. This is more intelligent than a continuous beep because it converts sensor readings into a decision.")
    heading(doc, "4.6 Analysis of Android application results", 2)
    paragraph(doc, "The Android app was developed to support both blind and deaf users. Onboarding separates the visual path from the audio path. The visual path supports captions, camera sign backup, glove connection and word suggestions. The audio path uses speech output and can be extended with voice command support where Android services are available.")
    paragraph(doc, "The user interface was refined to hide developer-only text such as raw ADC labels, fallback explanations and camera diagnostic cards. This is important because a final user should not have to understand hardware vocabulary before using the app. Technical details remain available in logs and firmware for debugging.")
    table(doc, "Table 22: Test case matrix.", ["Test case", "Input condition", "Expected behaviour", "Status"], [
        ("TC1", "Glove powered and BLE advertising", "Phone discovers SignaSenseGlove", "Prepared."),
        ("TC2", "Hold a valid letter steadily", "Letter enters word after stability time", "Implemented."),
        ("TC3", "Short noisy movement", "Letter is rejected", "Implemented."),
        ("TC4", "Camera mode selected", "Phone detects hand landmarks and predicts letter", "Implemented with model assets."),
        ("TC5", "No user-defined signs exist", "Training mode starts before custom mode", "Implemented."),
        ("TC6", "Obstacle enters warning range", "Stick speaks distance and action once", "Implemented in firmware logic."),
        ("TC7", "Very close obstacle", "Stick says stop now and starts left/right scan", "Implemented in guidance logic."),
    ], font_size=8.5)
    add_figure(doc, FIG_DIR / "figure_16_calibration_acceptance.png", "Figure 5: Calibration and acceptance-testing workflow.")
    doc.add_page_break()
    heading(doc, "CHAPTER FIVE: SYSTEM DEVELOPMENT AND IMPLEMENTATION", 1)
    heading(doc, "5.1. Requirements identified", 2)
    paragraph(doc, "Requirements were derived from the concept paper, project objectives, hardware testing and app interaction needs.")
    heading(doc, "5.1.1 User Requirements", 3)
    table(doc, "Table 5: User requirements.", ["User", "Requirement", "Reason"], [
        ("Blind user", "Receive spoken obstacle warnings", "Supports safe movement."),
        ("Blind user", "Use voice-guided onboarding", "Reduces dependence on visual controls."),
        ("Deaf user", "View captions and detected letters", "Supports visual communication."),
        ("Deaf user", "Build letters into words before speech", "Prevents wrong immediate speech."),
        ("Presenter/user", "Use camera backup", "Keeps sign detection available when glove battery is low."),
        ("All users", "Use dark or light mode", "Improves readability in different environments."),
    ], font_size=9)
    heading(doc, "5.1.2 Functional Requirements", 3)
    table(doc, "Table 6: Functional requirements.", ["ID", "Requirement", "Implementation"], [
        ("FR1", "Read five flex sensors continuously", "ESP32 analogRead with smoothing and calibration."),
        ("FR2", "Send glove data to phone", "BLE service named SignaSenseGlove."),
        ("FR3", "Detect and suggest words", "Android app letter filtering and dictionary suggestions."),
        ("FR4", "Detect obstacle distance", "Ultrasonic trigger/echo with filtered samples."),
        ("FR5", "Speak guidance", "Bluetooth speaker output and backup buzzer."),
        ("FR6", "Detect signs with camera backup", "CameraX plus MediaPipe hand landmarks."),
        ("FR7", "Train local custom signs", "Private device storage for user-defined samples."),
    ], font_size=9)
    heading(doc, "5.2 System requirements (of the new system)", 2)
    table(doc, "Table 7: System requirements.", ["Category", "Requirement"], [
        ("Hardware", "ESP32 board, five flex sensors, resistors, ultrasonic sensor, buzzer, Bluetooth speaker, Android phone."),
        ("Software", "Arduino IDE, Kotlin Android app, Android SDK, CameraX, MediaPipe Tasks Vision, Android TTS and BLE APIs."),
        ("Power", "Stable 3.3 V sensor supply, common ground, USB or power-bank supply for ESP32 boards."),
        ("Safety", "Voltage divider or level shifter for ultrasonic ECHO if the sensor outputs 5 V."),
        ("Connectivity", "BLE between glove and phone; Bluetooth audio between stick and speaker."),
    ], font_size=9)
    heading(doc, "5.3 Architectural design of the new system", 2)
    heading(doc, "5.3.1 Description of the Designed System", 3)
    paragraph(doc, "SignaSense is modular. The glove is responsible for sensing finger bends and sending communication data. The phone is responsible for interaction, captions, speech and camera backup. The stick is responsible for obstacle distance, local guidance and direct Bluetooth speaker output.")
    add_figure(doc, FIG_DIR / "figure_02_system_architecture.png", "Figure 6: SignaSense integrated system architecture.")
    heading(doc, "5.3.2 Hardware Specifications", 3)
    table(doc, "Table 8: Smart glove hardware pin allocation.", ["Finger", "ESP32 GPIO", "Signal"], [
        ("Thumb", "GPIO25", "Flex sensor voltage-divider output"),
        ("Index", "GPIO33", "Flex sensor voltage-divider output"),
        ("Middle", "GPIO32", "Flex sensor voltage-divider output"),
        ("Ring", "GPIO35", "Flex sensor voltage-divider output"),
        ("Pinky", "GPIO34", "Flex sensor voltage-divider output"),
    ], font_size=9)
    add_figure(doc, figures[3], "Figure 7: Smart glove flex-sensor voltage-divider wiring.")
    add_photo_placeholder(doc, "Figure 8: Final smart glove prototype photograph placeholder.", "NOTE: Insert a clear photograph of the finished glove showing the ESP32, five flex sensors and wiring.")
    table(doc, "Table 9: Smart stick hardware pin allocation.", ["Component", "ESP32 GPIO", "Note"], [
        ("Ultrasonic TRIG", "GPIO5", "Output pulse from ESP32"),
        ("Ultrasonic ECHO", "GPIO18", "Must be level-shifted/voltage-divided if sensor echo is 5 V"),
        ("Buzzer", "GPIO15", "Backup alert when speaker is unavailable"),
    ], font_size=9)
    add_figure(doc, figures[5], "Figure 9: Smart stick ultrasonic sensor and buzzer wiring.")
    add_photo_placeholder(doc, "Figure 10: Final smart stick prototype photograph placeholder.", "NOTE: Insert a clear photograph of the final smart stick showing ultrasonic sensor, ESP32, buzzer and power arrangement.")
    heading(doc, "5.3.3 Software Specifications", 3)
    table(doc, "Table 10: Software tools and libraries used.", ["Tool/Library", "Use in project"], [
        ("Arduino IDE", "ESP32 glove and stick firmware development."),
        ("Kotlin / Android SDK", "SignaSense Android application."),
        ("Android BLE APIs", "Connection between phone and smart glove."),
        ("Android TextToSpeech", "Speech output for app prompts and committed words."),
        ("CameraX ImageAnalysis", "Camera frames for sign backup."),
        ("MediaPipe Hand Landmarker", "Hand landmark detection for camera signs."),
        ("ESP32 A2DP Source", "Bluetooth speaker output for the smart stick."),
    ], font_size=9)
    heading(doc, "5.3.4 Systems Architecture", 3)
    paragraph(doc, "The system architecture separates sensing, processing and output. The glove sends BLE JSON status to the phone; the camera pipeline processes phone frames locally; the smart stick processes distance locally and outputs speech directly to the Bluetooth speaker.")
    add_figure(doc, FIG_DIR / "figure_09_camera_pipeline.png", "Figure 11: Camera sign backup recognition pipeline.")
    add_figure(doc, FIG_DIR / "figure_10_user_defined_training.png", "Figure 12: User-defined sign training workflow.")
    add_figure(doc, FIG_DIR / "figure_14_ble_data_exchange.png", "Figure 13: BLE data exchange between the smart glove and Android app.")
    add_figure(doc, FIG_DIR / "figure_15_deployment_architecture.png", "Figure 14: SignaSense deployment architecture.")
    heading(doc, "5.4 Module design and implementation", 2)
    heading(doc, "5.4.1 Smart glove firmware module", 3)
    paragraph(doc, "The glove firmware reads the five flex sensors continuously, smooths the readings and sends the current state through BLE. The firmware is intentionally modular so that each finger can be tested independently. Pin names are kept descriptive because the physical wiring changed several times during construction and the final report must preserve the final pin status.")
    paragraph(doc, "The glove does not try to guess a full word inside the ESP32. This responsibility is moved to the Android app because the phone has more memory, better text handling, TextToSpeech support and a user interface for accepting or clearing suggested words. The ESP32 therefore works as a sensor node while the phone works as the interpretation node.")
    table(doc, "Table 20: Glove BLE payload fields.", ["Field", "Example", "Purpose"], [
        ("device", "SignaSenseGlove", "Identifies the glove service."),
        ("raw", "[120, 340, 410, 95, 280]", "Carries thumb, index, middle, ring and pinky readings."),
        ("states", "[STRAIGHT, BENT, HALF, STRAIGHT, BENT]", "Carries classified finger states."),
        ("letter", "Y", "Carries the current stable detected letter when available."),
        ("word", "YES", "Carries the current committed or building word."),
        ("status", "Connected", "Shows whether the glove is active and readable."),
    ], font_size=8.5)
    heading(doc, "5.4.2 Android glove interpretation module", 3)
    paragraph(doc, "The Android glove module receives BLE status, applies additional filtering and updates the word builder. The app uses a stability rule so that a detected letter must remain present long enough before it is treated as intentional. This reduces the effect of electrical noise, loose jumper wires and accidental hand movement.")
    paragraph(doc, "The app also provides word suggestions from the letters that have already been collected. A user can accept a suggested word or clear the current word. Speech output is produced only after a word or caption is committed, which reduces wrong speech during presentation and real use.")
    heading(doc, "5.4.3 Camera sign backup module", 3)
    paragraph(doc, "Camera mode is included as a backup when the glove battery is low or when the user wants to demonstrate sign detection without the sensor glove. It uses the phone camera and model assets stored in the project. The model should be used in a controlled environment with the hand inside the camera view, enough light and a simple background.")
    table(doc, "Table 21: Camera backup operating conditions.", ["Condition", "Recommended setting", "Reason"], [
        ("Lighting", "Bright indoor or daylight", "Improves hand landmark visibility."),
        ("Background", "Plain background", "Reduces false hand segmentation."),
        ("Hand position", "Full hand inside frame", "Allows all fingers to be analysed."),
        ("Motion signs", "Hold start and end positions clearly", "Dynamic gestures need stable frames."),
        ("User-defined signs", "Train each sign with repeated samples", "Improves personal recognition."),
    ], font_size=8.5)
    heading(doc, "5.4.4 Smart stick firmware module", 3)
    paragraph(doc, "The smart stick firmware measures distance in centimetres and metres, filters invalid readings and avoids overlapping speech. The speech queue is important because the user should hear complete phrases. If a distance phrase is already playing, a new phrase waits unless the situation has become dangerous enough to require an immediate stop warning.")
    paragraph(doc, "The smart guide logic uses threshold levels: clear path, warning range, danger range and immediate stop range. In the warning range, the system gives distance and direction advice. In the danger range, it says obstacle ahead and prepares for a left/right scan. In the immediate stop range, it says stop now before additional instructions.")
    heading(doc, "5.4.5 Bluetooth speaker output module", 3)
    paragraph(doc, "The stick uses the ESP32 built-in Bluetooth capability to output audio to a nearby Bluetooth speaker. The selected speaker name is SINOBAND Book. In the final design, the buzzer remains as backup because Bluetooth audio can fail if the speaker battery is low, if the speaker is already paired to another phone, or if the user moves out of range.")
    heading(doc, "5.4.6 User interface and accessibility module", 3)
    paragraph(doc, "The app interface was redesigned to use a small, clean layout with dark and light mode options. The final interface focuses on actions that matter to the user: connect glove, camera backup, user-defined signs, current word, suggestions, speak and clear. Developer details such as ADC terminology and fallback logic are not shown on the main screen.")
    paragraph(doc, "For blind users, the audio interface should guide the user through language selection and navigation. For deaf users, the visual interface should show clear captions and avoid unnecessary technical text. The design therefore separates presentation mode from debugging mode.")
    heading(doc, "5.5.5 Data Requirements", 3)
    paragraph(doc, "The glove requires raw flex readings, bend percentages, detected letter, current word and connection status. The stick requires pulse duration, calibrated distance, warning state, scan phase and guidance phrase. The app requires language choice, UI mode, current letter, current word, sentence suggestions and local custom sign samples.")
    heading(doc, "5.5.6 Data Inputs (database and tables)", 3)
    paragraph(doc, "The prototype does not use an online database. Inputs are local: analogue sensor readings, BLE JSON messages, button selections, camera frames and user-defined sign samples stored privately on the phone.")
    heading(doc, "5.5.7 System Flow", 3)
    paragraph(doc, "The glove flow is: read sensors, smooth readings, calibrate, classify finger states, detect stable letter, send status through BLE, build word in app and speak only after user confirmation.")
    paragraph(doc, "The stick flow is: measure distance, reject invalid readings, filter samples, speak obstacle warning only inside range, ask the user to scan left and right when danger is close, compare both sides and speak the safest movement suggestion.")
    add_figure(doc, figures[7], "Figure 15: Smart stick left/right scanning procedure for route guidance.")
    add_figure(doc, FIG_DIR / "figure_11_stick_algorithm.png", "Figure 16: Smart stick distance filtering and guidance logic.")
    heading(doc, "5.5.8 Physical database/system design", 3)
    paragraph(doc, "Physical data storage is minimal. The app stores user-defined camera sign samples only on the installed phone. This protects the general model from being changed by one user's custom alphabet and supports users who cannot make standard signs.")
    heading(doc, "5.5.9 Data Outputs (System forms). These should be very well explained. (Remember this is Presentation and Discussion of results)", 3)
    paragraph(doc, "The main outputs are app captions, current letters, word suggestions, sentence suggestions, spoken committed words, spoken obstacle distance, spoken movement guidance and backup buzzer alerts. The app UI is designed to show user-facing outputs only, while developer diagnostics remain inside the code or serial monitor.")
    if (FIG_DIR / "figure_app_dashboard.png").exists():
        add_figure(doc, FIG_DIR / "figure_app_dashboard.png", "Figure 17: SignaSense Android dashboard and accessibility controls.", width=3.0)


def add_conclusion_and_appendices(doc, figures):
    doc.add_page_break()
    heading(doc, "CHAPTER SIX: SUMMARY, CONCLUSION AND RECOMMENDATIONS", 1)
    heading(doc, "6.1 Summary (optional)", 2)
    paragraph(doc, "SignaSense was developed as an integrated assistive prototype for communication and mobility support. It includes a smart glove, smart stick and Android app. The glove supports letters and word building, the stick supports distance and guidance, and the app supports captions, speech, onboarding, camera backup and local user-defined signs.")
    heading(doc, "6.2 Conclusion", 2)
    paragraph(doc, "The project achieved its core aim of designing a low-cost assistive prototype that combines communication and mobility support. It demonstrates that ESP32 boards, flex sensors, an ultrasonic sensor, Bluetooth output and an Android app can be integrated into a practical final-year prototype.")
    paragraph(doc, "6.2.1 Legal and Ethical concerns")
    paragraph(doc, "The project involves persons with disabilities and therefore requires informed consent, privacy, respectful testing and safe supervision. Camera and voice features should be explained to users. User-defined sign samples must remain on the user's phone unless the user gives permission to share them.")
    paragraph(doc, "6.2.2 Also, the researcher can conclude about the positive implications of adopting the new system to the case study and how the rest of the society can benefit")
    paragraph(doc, "Adopting SignaSense can improve independent communication, route awareness and local technical capacity for assistive device innovation.")
    paragraph(doc, "6.2.3 Were the objectives realized in relation to the user requirements?")
    paragraph(doc, "The objectives were realised at prototype level. The app installs, the camera model loads, BLE glove communication works, and stick firmware was designed for spoken obstacle guidance. Further hardware stabilisation is required for consistent flex-sensor readings.")
    paragraph(doc, "6.2.4 What are the other factors that influence the proper adoption of the new systems?")
    paragraph(doc, "Adoption depends on power stability, durable wiring, clear enclosure design, user training, local-language prompt quality, affordability and supervised field testing.")
    heading(doc, "6.3 Recommendations", 2)
    bullet(doc, "Move glove wiring from loose jumper wires to soldered perfboard or PCB connections.")
    bullet(doc, "Use fixed resistor values for every flex sensor voltage divider and document calibration readings.")
    bullet(doc, "Add more sensors to the stick for left, right and ground-level detection without manual sweeping.")
    bullet(doc, "Improve local-language speech prompts for English, Luganda and Acholi.")
    bullet(doc, "Conduct supervised testing with deaf and blind participants and record accuracy, response time and comfort.")
    heading(doc, "6.4 Areas for further study", 2)
    bullet(doc, "Machine-learning recognition for a larger sign-language vocabulary and dynamic gestures.")
    bullet(doc, "Offline voice command recognition without extra hardware cost.")
    bullet(doc, "Battery-life optimisation and safer enclosure design.")
    bullet(doc, "Haptic feedback for users who may not hear the phone or speaker.")
    bullet(doc, "Full multilingual localisation and field testing in community settings.")
    doc.add_page_break()
    heading(doc, "REFERENCES", 1)
    for ref in [
        "Android Developers (2026a) SpeechRecognizer. Available at: https://developer.android.com/reference/android/speech/SpeechRecognizer.html (Accessed: 4 May 2026).",
        "Android Developers (2026b) TextToSpeech. Available at: https://developer.android.com/reference/android/speech/tts/TextToSpeech (Accessed: 4 May 2026).",
        "Android Developers (2026c) Image analysis. Available at: https://developer.android.com/training/camerax/analyze (Accessed: 4 May 2026).",
        "Espressif Systems (2026) ESP-FAQ Handbook. Available at: https://documentation.espressif.com/api/resource/path/docs%2Fprojects%2Fesp-faq%2Fen%2Flatest%2Fesp-faq-en-master.pdf (Accessed: 4 May 2026).",
        "Gallaudet University (2018) K-12 ASL Content Standards. Available at: https://gallaudet.edu/wp-content/uploads/gcloud/gal-media/Documents/ASL-Standards/K-12-ASL-Content-Standards.pdf (Accessed: 4 May 2026).",
        "Google AI Edge (2025) Hand landmarks detection guide for Android. Available at: https://ai.google.dev/edge/mediapipe/solutions/vision/hand_landmarker/android (Accessed: 4 May 2026).",
        "Uganda (2020) Persons with Disabilities Act, 2020. Available at: https://media.ulii.org/media/legislation/118805/source_file/6c59741d41db2532/2020-3.pdf (Accessed: 4 May 2026).",
        "World Health Organization (2026a) Deafness and hearing loss. Available at: https://www.who.int/news-room/fact-sheets/detail/deafness-and-hearing-loss (Accessed: 4 May 2026).",
        "World Health Organization (2026b) Blindness and vision impairment. Available at: https://www.who.int/news-room/fact-sheets/detail/blindness-and-visual-impairment (Accessed: 4 May 2026).",
    ]:
        paragraph(doc, ref)
    doc.add_page_break()
    heading(doc, "APPENDICES", 1)
    paragraph(doc, "The appendices provide assessment, typographic, supervision, workplan and face-page materials required by the undergraduate report format.")
    heading(doc, "INDICES", 1)
    paragraph(doc, "Index terms: accessibility, Android application, assistive technology, Bluetooth Low Energy, camera sign backup, ESP32, flex sensor, smart glove, smart stick, ultrasonic sensor, user-defined signs.")
    doc.add_page_break()
    heading(doc, "APPENDIX A: ASSESMENT FORM", 1)
    table(doc, "Assessment form summary", ["Criteria", "Maximum Score", "Awarded Score"], [
        ("Problem statement, background, aims and objectives", "5", ""),
        ("Literature review", "3", ""),
        ("Methodology", "5", ""),
        ("System development and implementation", "20", ""),
        ("Testing, results and discussion", "20", ""),
        ("Report formatting and presentation", "10", ""),
        ("Oral presentation and demonstration", "25", ""),
    ], font_size=9)
    paragraph(doc, "Assessment note: the prototype should be assessed against both technical operation and social usefulness. Technical assessment checks whether the glove sends data, whether the app builds words, whether the camera backup loads and whether the smart stick speaks meaningful guidance. Social usefulness checks whether the outputs can be understood by intended users.")
    table(doc, "Table 23: User acceptance checklist.", ["Item", "Acceptance question", "Result"], [
        ("Glove comfort", "Can the glove be worn without pain or restricted circulation?", ""),
        ("Letter display", "Can the user see the detected letter clearly?", ""),
        ("Word confirmation", "Can the user correct a wrong letter before speech?", ""),
        ("Audio clarity", "Can the spoken output be understood in a normal room?", ""),
        ("Stick warning", "Does the stick warn before the obstacle is too close?", ""),
        ("Guidance", "Can the user follow the left/right scan instruction?", ""),
        ("Camera backup", "Can camera mode work when the glove is not powered?", ""),
        ("Theme", "Can the user choose a readable dark or light interface?", ""),
    ], font_size=8.5)
    heading(doc, "Appendix A.1 Interview guide", 2)
    paragraph(doc, "The interview guide below can be used during supervised testing. It should be administered respectfully and only after the participant has understood that SignaSense is a prototype.")
    for item in [
        "What communication or mobility problem affects you most in daily activities?",
        "Was the glove comfortable enough for short use?",
        "Were the captions and suggested words easy to read?",
        "Was the spoken output clear enough to understand?",
        "Did the stick warning come early enough for safe reaction?",
        "Which language would you prefer for speech prompts?",
        "What change would make the system easier to use?",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "APPENDIX B: TYPOGRAPHIC FORMAT", 1)
    paragraph(doc, "The report uses Times New Roman, font size 12 for body text, 1.5 line spacing, justified paragraphs where appropriate, and formal heading levels. Tables and captions are kept readable, and figures are numbered in the order they appear.")
    paragraph(doc, "Formatting applied in this generated report includes portrait orientation, standard margins, title page, declaration, approval page, acknowledgement, abstract, table of contents, lists of figures and tables, Harvard-style references and appendices. Tables use compact font sizes where necessary so that wiring and test information remains readable.")
    paragraph(doc, "The final printed copy should use the same document without changing the heading order. After inserting the physical prototype photographs in the marked placeholders, the list of figures and table of contents should be refreshed so that page numbers match the final printed version.")
    heading(doc, "APPENDIX C: SUPERVISION MONITORING FORM", 1)
    table(doc, "Supervision monitoring form", ["Date of Meeting", "Objective", "Agreed action", "Students", "Supervisor comment"], [
        ("", "Project topic and scope", "", "", ""),
        ("", "Hardware wiring review", "", "", ""),
        ("", "App and firmware review", "", "", ""),
        ("", "Testing and report review", "", "", ""),
    ], font_size=8.5)
    heading(doc, "Appendix C.1 Prototype issue log", 2)
    table(doc, "Prototype issue log", ["Issue", "Likely cause", "Corrective action"], [
        ("Only one flex sensor reads", "Other sensors not connected as voltage dividers", "Add resistor-to-ground for every sensor and confirm common GND."),
        ("All values are zero", "No valid analogue voltage reaching ESP32", "Check 3.3 V, centre tap and ground continuity."),
        ("Random letters appear", "Noise, floating pins or unstable power", "Use stability timing and improve wiring."),
        ("Bluetooth speaker has no voice", "Speaker not paired or already connected elsewhere", "Pair to SINOBAND Book and keep it near the stick."),
        ("Distance repeats too much", "Speech interval too short", "Use phrase queue and speak only on meaningful changes."),
        ("Camera predictions are weak", "Lighting, angle or model mismatch", "Use camera backup in good lighting and train custom samples."),
    ], font_size=8.2)
    doc.add_page_break()
    heading(doc, "APPENDIX D: RESEACH WORKPLAN", 1)
    table(doc, "Table 15: Research work plan.", ["Week(s)", "Activity", "Output"], [
        ("1-2", "Topic selection and concept refinement", "Approved project idea"),
        ("2-5", "Literature review and proposal writing", "Proposal draft"),
        ("5-7", "Hardware design and wiring", "Glove/stick circuits"),
        ("7-11", "Firmware development", "Arduino sketches"),
        ("8-13", "Android app development", "SignaSense APK"),
        ("13-15", "Integration and testing", "Working prototype"),
        ("15-16", "Report writing and presentation preparation", "Final report"),
    ], font_size=9)
    add_figure(doc, FIG_DIR / "figure_13_workplan_gantt.png", "Figure 18: SignaSense project work plan Gantt chart.")
    heading(doc, "Appendix D.1 Detailed work plan narrative", 2)
    paragraph(doc, "The first stage covered topic identification, problem analysis and approval of the concept. The project was selected because it combines two visible community needs: communication support for deaf users and mobility support for blind users. The concept stage also defined the scope so that the project could be completed using available ESP32 boards, flex sensors, an ultrasonic sensor and an Android phone.")
    paragraph(doc, "The second stage covered literature review and proposal preparation. This stage clarified that the glove alone cannot solve complete sign-language grammar, and the stick alone cannot support communication. The reviewed literature therefore supported an integrated approach and justified the inclusion of camera backup and user-defined signs.")
    paragraph(doc, "The third stage covered hardware construction. The glove wiring focused on five flex-sensor voltage dividers. The stick wiring focused on ultrasonic trigger/echo, buzzer backup and Bluetooth speaker output. Several pin mappings were tested before settling on the final glove pins documented in Chapter Five.")
    paragraph(doc, "The fourth stage covered firmware and app development. The glove firmware was designed to send readings to the phone. The stick firmware was refined to avoid repeated phrases and to guide the user through left/right scanning. The Android app was renamed SignaSense, given an icon, dark and light mode, camera backup and user-defined sign training.")
    paragraph(doc, "The final stage covered integration, testing and report writing. Rendering the report as page images was used to confirm that headings, diagrams, tables and appendices were visible and not clipped.")
    doc.add_page_break()
    heading(doc, "APPENDIX E: FACE PAGE", 1)
    paragraph(doc, "[PROJECT TITLE]: SIGNASENSE - AN INTEGRATED ASSISTIVE SOLUTION FOR THE DEAF AND BLIND", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc, "DEPARTMENT OF INFORMATION TECHNOLOGY, INSTITUTE OF COMPUTER SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc, f"Supervisor: {SUPERVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Student names and registration numbers are provided on the report face page and declaration page.", align=WD_ALIGN_PARAGRAPH.CENTER)
    heading(doc, "Appendix E.1 Prototype demonstration guide", 2)
    paragraph(doc, "During presentation, the demonstrator should begin by showing the SignaSense app onboarding, then the deaf visual interface, then the blind audio interface. The smart glove demonstration should show how detected letters enter the word area and how the user can speak a committed word. The camera backup should be presented as a fallback when glove battery or wiring is unavailable.")
    paragraph(doc, "The smart stick demonstration should be done in a controlled space. Place an obstacle ahead, allow the stick to speak the distance and warning, then perform the left and right scan slowly. The presenter should explain that the system compares the two sides and chooses the side with more free distance.")
    paragraph(doc, "The presenter should not claim that the prototype is a certified medical device. It is a final-year assistive technology prototype that demonstrates feasibility, integration and local innovation.")
    doc.add_page_break()
    heading(doc, "Appendix E.2 Android application user guide", 2)
    paragraph(doc, "The Android application is the user-facing centre of SignaSense. A user first chooses the appropriate path during onboarding. A deaf user continues to the visual interface, where glove letters, camera signs, current word and suggestions are shown. A blind user continues to the audio interface, where speech prompts guide the interaction.")
    paragraph(doc, "In smart glove mode, the user connects to the glove through Bluetooth Low Energy. After connection, the app receives sensor state updates and displays useful results. The user should hold a sign steadily until it is accepted into the word area. If the wrong letter appears, the clear control can be used before speaking the word.")
    paragraph(doc, "In camera backup mode, the user places the hand in front of the phone camera. The mode is intended for backup use when glove power is low or when the presenter wants to show a non-contact sign recognition option. The camera should be used in good lighting with the full hand visible.")
    paragraph(doc, "In user-defined signs mode, the phone stores a personal alphabet on that specific device. The trained signs do not affect other phones or the general model. This protects the general app while still allowing users with different hand abilities to define signs they can make comfortably.")
    for item in [
        "Use dark mode in low light or during night demonstrations.",
        "Use light mode in bright rooms or outdoor conditions.",
        "Speak a word only after confirming that the built word is correct.",
        "Retrain user-defined signs if camera predictions become inconsistent.",
        "Keep the phone camera lens clean during camera backup use.",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "Appendix E.3 Smart glove setup guide", 2)
    paragraph(doc, "The smart glove should be assembled so that every flex sensor has two electrical paths: one path to 3.3 V through the flex sensor and one path to ground through a fixed resistor. The ESP32 analogue pin reads from the centre point of this voltage divider. Without the resistor, the analogue pin can float and the app may show random letters or no bend response.")
    paragraph(doc, "The recommended setup sequence is to wire one finger first, upload the raw-reading test firmware, bend the sensor and confirm that the value changes. After one finger works, repeat the same wiring method for the remaining fingers. This prevents a situation where all five sensors are wired incorrectly at once and the cause becomes difficult to isolate.")
    paragraph(doc, "The final glove pin status used in this report is thumb GPIO25, index GPIO33, middle GPIO32, ring GPIO35 and pinky GPIO34. GPIO34 and GPIO35 are input-only pins on ESP32, which is acceptable for analogue sensor input. All sensor grounds must be connected to the ESP32 ground.")
    table(doc, "Glove setup checklist", ["Check", "Expected result"], [
        ("ESP32 powered", "Board LED or serial output confirms power."),
        ("Common ground", "All sensor grounds join the ESP32 GND."),
        ("Voltage divider", "Every flex sensor has a resistor to ground."),
        ("Raw readings", "Each pin gives a value above zero and below the ADC maximum."),
        ("Bend response", "The active finger changes when bent."),
        ("BLE advertising", "Phone can scan and find SignaSenseGlove."),
    ], font_size=8.5)
    doc.add_page_break()
    heading(doc, "Appendix E.4 Smart stick setup and field guide", 2)
    paragraph(doc, "The smart stick should be tested indoors before outdoor use. The ultrasonic sensor must face forward and should not be blocked by the stick body or loose wires. If the ultrasonic sensor outputs a 5 V echo signal, the echo line must pass through a voltage divider or level shifter before reaching GPIO18 because ESP32 pins are not 5 V tolerant.")
    paragraph(doc, "The speaker should be switched on and kept near the ESP32 during connection. The expected Bluetooth speaker name is SINOBAND Book. If the speaker does not produce voice, check whether it is already connected to another device, whether its battery is charged and whether the stick firmware is still using Bluetooth speaker output rather than app Wi-Fi mode.")
    paragraph(doc, "The user should not walk immediately after hearing the first warning. In danger mode, the stick asks for left and right scans. The holder should slowly point or sweep the stick left when asked, then right when asked. After both sides are sampled, the stick speaks which side has more free distance.")
    for item in [
        "Test 30 cm, 50 cm and 100 cm distances using a tape measure.",
        "Do not test near stairs, roads or crowded spaces during early trials.",
        "Keep the backup buzzer connected in case Bluetooth audio fails.",
        "Use a power bank that can supply enough current for the ESP32 and Bluetooth activity.",
        "Place sensor wiring inside an enclosure before field trials.",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "Appendix E.5 Presentation demonstration script", 2)
    paragraph(doc, "The demonstration should be structured so that the panel sees the problem, the solution and the tested behaviour. First, introduce SignaSense as an integrated assistive solution for deaf and blind users. Then explain that the project has three parts: a smart glove for sign input, a smart stick for obstacle guidance and an Android app for captions, suggestions, speech and camera backup.")
    paragraph(doc, "For the glove demonstration, wear the glove, connect it from the app and show a few stable letters. Explain that random movement is filtered so that letters only enter the word area after being held steadily. Then commit the word and allow the phone to speak it. If the hardware readings are unstable, switch to camera backup and explain that it is included for continuity.")
    paragraph(doc, "For the camera backup demonstration, open camera mode and make a visible sign in front of the phone. Explain that the camera mode can work with a glove or bare hand, but in the project it supports the glove presentation when battery or wiring limits occur. Then show the user-defined signs option and explain how a user can train personal signs on only their own device.")
    paragraph(doc, "For the smart stick demonstration, place an obstacle ahead and allow the stick to speak the obstacle warning. Move the stick closer so that it says stop now. Then perform the left and right scan slowly and let the system recommend the better side. End by explaining that the buzzer remains as backup if the Bluetooth speaker is not connected.")
    doc.add_page_break()
    heading(doc, "Appendix E.6 Maintenance and improvement notes", 2)
    paragraph(doc, "The prototype should be maintained by checking wires, sensor placement, power supply and app version before every demonstration. Loose jumper wires were one of the biggest causes of unstable readings during development. A soldered board or printed circuit board would make the system more reliable.")
    paragraph(doc, "The glove can be improved by adding an inertial measurement unit to detect palm orientation and motion. This would help letters such as J and Z, which require movement, and signs that look similar from finger bend alone. The camera backup can also be improved by collecting more local training samples under different lighting conditions.")
    paragraph(doc, "The smart stick can be improved by adding side-facing sensors instead of requiring manual left/right scanning. It can also include vibration feedback for noisy environments and a stronger enclosure to protect the ESP32 and ultrasonic module. The app can be improved with more Luganda and Acholi prompts and better offline language support.")
    paragraph(doc, "These improvements are recommended for future work after the submitted prototype has been demonstrated, assessed and documented.")
    doc.add_page_break()
    heading(doc, "Appendix E.7 Final submission checklist", 2)
    paragraph(doc, "This checklist is included to help the group prepare the final submitted copy and demonstration materials. It keeps the report aligned with the physical prototype and reduces the chance of submitting a document that is missing photographs, signatures or calibration evidence.")
    for item in [
        "Insert the final smart glove photograph at Figure 8 before printing.",
        "Insert the final smart stick photograph at Figure 10 before printing.",
        "Confirm that all group members sign the declaration page.",
        "Confirm that the supervisor signs the approval page after review.",
        "Carry the ESP32 boards, phone, Bluetooth speaker, charger and power bank for presentation.",
        "Prepare a short live demonstration for glove letters, camera backup and smart stick guidance.",
        "Carry spare jumper wires, resistors, USB cable and tape for quick repair during presentation.",
        "Check that the APK installed on the phone is named SignaSense and opens correctly.",
        "Confirm that the smart stick speaker target remains SINOBAND Book.",
        "Refresh page numbers and lists after inserting final physical photographs.",
    ]:
        bullet(doc, item)
    paragraph(doc, "After completing the checklist, the report can be printed and bound according to the department's submission instructions.")


def main():
    figures = generate_figures()
    figures.update(generate_extra_figures())
    doc = Document(TEMPLATE_DOCX)
    clear_document(doc)
    set_document_defaults(doc)
    add_title_page(doc)
    add_front_matter(doc)
    add_introduction(doc)
    add_concept_paper(doc, figures)
    add_project_proposal(doc, figures)
    add_literature_review(doc)
    add_methodology(doc)
    add_report_chapters(doc, figures)
    add_conclusion_and_appendices(doc, figures)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
