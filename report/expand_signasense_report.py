from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parent
BASE_DOCX = BASE_DIR / "SignaSense_Final_Report.docx"
OUT_DOCX = BASE_DIR / "SignaSense_Final_Report_Full.docx"
FIG_DIR = BASE_DIR / "generated_figures"
FIG_DIR.mkdir(exist_ok=True)


GROUP_MEMBERS = [
    ("NALUBOWA BENITAH MARGRET", "2023/BIT/169/PS"),
    ("ADEKE MARY", "2023/BIT/034/PS"),
    ("MUHUMUZA MUHAMMAD", "2023/BIT/252/PS"),
    ("MIREMBE JERMIMAH ARNN", "2023/BIT/148/PS"),
    ("KANSIIME PIUS", "2023/BIT/120/PS"),
]


BLUE = "#D9EAF7"
BLUE_DARK = "#1F4E79"
GREEN = "#DFF3E3"
YELLOW = "#FFF2CC"
GREY = "#F3F4F6"
INK = "#111827"
LINE = "#344054"


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, fnt, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_center_text(draw, xy, text, fnt, fill=INK, max_width=None, line_spacing=8):
    x1, y1, x2, y2 = xy
    lines = wrap_text(draw, text, fnt, (x2 - x1 - 30) if max_width is None else max_width)
    total_h = sum(draw.textbbox((0, 0), line, font=fnt)[3] for line in lines) + line_spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=fill, font=fnt)
        y += (bbox[3] - bbox[1]) + line_spacing


def box(draw, xy, text, fill=BLUE, outline=LINE, fnt=None, radius=18):
    if fnt is None:
        fnt = font(25, bold=True)
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    draw_center_text(draw, xy, text, fnt)


def arrow(draw, start, end, color=LINE, width=5):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 20 * direction, y2 - 12), (x2 - 20 * direction, y2 + 12)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 12, y2 - 20 * direction), (x2 + 12, y2 - 20 * direction)]
    draw.polygon(pts, fill=color)


def canvas(title):
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 95), fill=BLUE_DARK)
    draw.text((50, 28), title, fill="white", font=font(38, bold=True))
    return img, draw


def save_figure(name, title, nodes, arrows_list):
    img, draw = canvas(title)
    for item in nodes:
        box(draw, item["xy"], item["text"], fill=item.get("fill", BLUE), fnt=font(item.get("size", 24), bold=True))
    for start, end in arrows_list:
        arrow(draw, start, end)
    path = FIG_DIR / name
    img.save(path)
    return path


def generate_figures():
    figures = {}

    figures[1] = save_figure(
        "figure_01_conceptual_framework.png",
        "SignaSense Conceptual Framework",
        [
            {"xy": (90, 170, 440, 340), "text": "User Inputs\nFinger bends\nHand signs\nObstacle distance", "fill": GREEN},
            {"xy": (560, 170, 930, 340), "text": "Processing\nESP32 firmware\nAndroid app\nCamera landmarks", "fill": BLUE},
            {"xy": (1050, 170, 1460, 340), "text": "Decision Logic\nLetters\nWords\nDistance buckets\nPath guidance", "fill": YELLOW},
            {"xy": (610, 550, 1160, 745), "text": "Accessible Outputs\nCaptions, speech, buzzer backup,\nBluetooth speaker voice guidance", "fill": GREY},
        ],
        [((440, 255), (560, 255)), ((930, 255), (1050, 255)), ((1255, 340), (960, 550)), ((730, 340), (840, 550))],
    )

    figures[2] = save_figure(
        "figure_02_system_architecture.png",
        "SignaSense Integrated System Architecture",
        [
            {"xy": (80, 180, 430, 350), "text": "Smart Glove\n5 flex sensors\nESP32 BLE", "fill": GREEN},
            {"xy": (610, 180, 1030, 350), "text": "SignaSense Android App\nOnboarding, captions,\nTTS, suggestions", "fill": BLUE},
            {"xy": (1230, 180, 1640, 350), "text": "Phone Outputs\nVisual captions\nSpoken words\nVoice prompts", "fill": YELLOW},
            {"xy": (80, 610, 430, 780), "text": "Camera Backup\nMediaPipe landmarks\nA-Z classifier", "fill": GREEN},
            {"xy": (610, 610, 1030, 780), "text": "Smart Stick\nUltrasonic sensor\nESP32 A2DP source", "fill": GREEN},
            {"xy": (1230, 610, 1640, 780), "text": "Bluetooth Speaker\nSINOBAND Book\nDistance and guidance", "fill": YELLOW},
        ],
        [((430, 265), (610, 265)), ((1030, 265), (1230, 265)), ((430, 695), (610, 300)), ((1030, 695), (1230, 695))],
    )

    figures[3] = save_figure(
        "figure_03_glove_wiring.png",
        "Smart Glove Flex-Sensor Voltage Divider Wiring",
        [
            {"xy": (80, 160, 420, 320), "text": "3.3 V Rail", "fill": YELLOW},
            {"xy": (80, 720, 420, 880), "text": "Common GND", "fill": GREY},
            {"xy": (620, 150, 980, 320), "text": "Thumb GPIO25\nIndex GPIO33\nMiddle GPIO32", "fill": GREEN},
            {"xy": (620, 390, 980, 560), "text": "Ring GPIO35\nPinky GPIO34", "fill": GREEN},
            {"xy": (1210, 260, 1630, 520), "text": "ESP32\nanalogRead()\ncalibration\nBLE JSON output", "fill": BLUE},
        ],
        [((420, 240), (620, 240)), ((420, 800), (620, 500)), ((980, 240), (1210, 350)), ((980, 475), (1210, 430))],
    )

    figures[5] = save_figure(
        "figure_05_stick_wiring.png",
        "Smart Stick Ultrasonic Sensor and Buzzer Wiring",
        [
            {"xy": (90, 180, 480, 380), "text": "Ultrasonic Sensor\nTRIG -> GPIO5\nECHO -> divider -> GPIO18", "fill": GREEN},
            {"xy": (90, 600, 480, 780), "text": "Backup Buzzer\nSignal -> GPIO15\nGND -> ESP32 GND", "fill": YELLOW},
            {"xy": (720, 250, 1120, 540), "text": "ESP32 Smart Stick\nsample filtering\ncalibrated distance\nA2DP voice queue", "fill": BLUE},
            {"xy": (1350, 250, 1700, 540), "text": "Bluetooth Speaker\nSINOBAND Book\nspoken distance and guidance", "fill": YELLOW},
        ],
        [((480, 280), (720, 340)), ((480, 690), (720, 470)), ((1120, 395), (1350, 395))],
    )

    figures[7] = save_figure(
        "figure_07_stick_scan.png",
        "Smart Stick Left/Right Scanning Procedure",
        [
            {"xy": (80, 180, 390, 330), "text": "Obstacle very close\nStop now", "fill": YELLOW},
            {"xy": (520, 180, 830, 330), "text": "Prompt user\nScan left slowly", "fill": BLUE},
            {"xy": (960, 180, 1270, 330), "text": "Collect left\nbest clear distance", "fill": GREEN},
            {"xy": (80, 610, 390, 760), "text": "Prompt user\nScan right slowly", "fill": BLUE},
            {"xy": (520, 610, 830, 760), "text": "Collect right\nbest clear distance", "fill": GREEN},
            {"xy": (960, 610, 1270, 760), "text": "Compare sides\nleft, right or no clear side", "fill": YELLOW},
            {"xy": (1400, 390, 1690, 570), "text": "Speak guidance\nMove left\nMove right\nNo clear side", "fill": GREY},
        ],
        [((390, 255), (520, 255)), ((830, 255), (960, 255)), ((1115, 330), (235, 610)), ((390, 685), (520, 685)), ((830, 685), (960, 685)), ((1270, 685), (1460, 570))],
    )

    figures[8] = save_figure(
        "figure_08_app_navigation.png",
        "SignaSense Android Application Navigation",
        [
            {"xy": (90, 160, 430, 320), "text": "Start\nAsk: blind or deaf", "fill": BLUE},
            {"xy": (600, 160, 940, 320), "text": "Blind path\nvoice prompts\nlanguage choice", "fill": YELLOW},
            {"xy": (1110, 160, 1450, 320), "text": "Deaf path\nvisual captions\nlarge controls", "fill": GREEN},
            {"xy": (340, 600, 680, 770), "text": "Smart Glove BLE\nletters, words,\nspeak commit", "fill": GREEN},
            {"xy": (820, 600, 1160, 770), "text": "Camera Signs\ngeneral/user-defined\nA-Z backup", "fill": GREEN},
            {"xy": (1300, 600, 1640, 770), "text": "Theme + Settings\ndark/light\nrepeat prompts", "fill": GREY},
        ],
        [((430, 240), (600, 240)), ((430, 240), (1110, 240)), ((770, 320), (510, 600)), ((1280, 320), (990, 600)), ((1450, 240), (1470, 600))],
    )

    figures[9] = save_figure(
        "figure_09_camera_pipeline.png",
        "Camera Sign Backup Recognition Pipeline",
        [
            {"xy": (80, 190, 380, 340), "text": "Phone camera\nvideo frame", "fill": GREEN},
            {"xy": (520, 190, 850, 340), "text": "MediaPipe\n21 hand landmarks", "fill": BLUE},
            {"xy": (990, 190, 1320, 340), "text": "Classifier\ngeometry + samples", "fill": YELLOW},
            {"xy": (1450, 190, 1730, 340), "text": "Letter A-Z", "fill": GREEN},
            {"xy": (520, 600, 850, 760), "text": "Hold filter\n30 ms stable\nrelease reset", "fill": GREY},
            {"xy": (990, 600, 1320, 760), "text": "Word builder\nsuggestions\nsentence options", "fill": BLUE},
        ],
        [((380, 265), (520, 265)), ((850, 265), (990, 265)), ((1320, 265), (1450, 265)), ((1590, 340), (690, 600)), ((850, 680), (990, 680))],
    )

    figures[10] = save_figure(
        "figure_10_user_defined_training.png",
        "User-Defined Sign Training on One Phone",
        [
            {"xy": (80, 180, 390, 330), "text": "Select\nUser-defined signs", "fill": BLUE},
            {"xy": (520, 180, 830, 330), "text": "Choose letter\nA to Z", "fill": GREEN},
            {"xy": (960, 180, 1270, 330), "text": "Capture 8 steady\nlandmark samples", "fill": YELLOW},
            {"xy": (1400, 180, 1690, 330), "text": "Save locally\nprivate app storage", "fill": GREY},
            {"xy": (520, 610, 830, 760), "text": "Complete 26 letters", "fill": GREEN},
            {"xy": (960, 610, 1270, 760), "text": "Recognise only\nthis user's signs", "fill": BLUE},
        ],
        [((390, 255), (520, 255)), ((830, 255), (960, 255)), ((1270, 255), (1400, 255)), ((1545, 330), (675, 610)), ((830, 685), (960, 685))],
    )

    figures[11] = save_figure(
        "figure_11_stick_algorithm.png",
        "Smart Stick Distance Filtering and Guidance Logic",
        [
            {"xy": (70, 165, 360, 310), "text": "Trigger ultrasonic\nsensor", "fill": GREEN},
            {"xy": (500, 165, 790, 310), "text": "Collect 7 samples", "fill": BLUE},
            {"xy": (930, 165, 1220, 310), "text": "Reject invalid\nor outlier echoes", "fill": YELLOW},
            {"xy": (1360, 165, 1680, 310), "text": "Calibrate\ncm and metres", "fill": GREEN},
            {"xy": (500, 610, 790, 760), "text": "Classify bucket\nwarning, danger,\nvery close", "fill": BLUE},
            {"xy": (930, 610, 1220, 760), "text": "Queue speech\nno overlap", "fill": YELLOW},
            {"xy": (1360, 610, 1680, 760), "text": "If very close\nstart scan cycle", "fill": GREY},
        ],
        [((360, 238), (500, 238)), ((790, 238), (930, 238)), ((1220, 238), (1360, 238)), ((1520, 310), (645, 610)), ((790, 685), (930, 685)), ((1220, 685), (1360, 685))],
    )

    figures[12] = save_figure(
        "figure_12_testing_workflow.png",
        "Prototype Verification and Testing Workflow",
        [
            {"xy": (80, 160, 410, 320), "text": "Hardware check\npower, GND,\npin continuity", "fill": GREEN},
            {"xy": (560, 160, 890, 320), "text": "Serial diagnostics\nraw values\ndistance logs", "fill": BLUE},
            {"xy": (1040, 160, 1370, 320), "text": "Connectivity\nBLE, A2DP,\nphone install", "fill": YELLOW},
            {"xy": (560, 600, 890, 760), "text": "Functional tests\nletters, words,\nobstacles", "fill": BLUE},
            {"xy": (1040, 600, 1370, 760), "text": "Record issues\ncalibrate\niterate", "fill": GREY},
        ],
        [((410, 240), (560, 240)), ((890, 240), (1040, 240)), ((1205, 320), (725, 600)), ((890, 680), (1040, 680))],
    )

    figures[13] = draw_workplan()
    return figures


def draw_workplan():
    img = Image.new("RGB", (1900, 1150), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1900, 95), fill=BLUE_DARK)
    draw.text((50, 28), "SignaSense Project Work Plan Gantt Chart", fill="white", font=font(38, bold=True))
    tasks = [
        ("Problem identification", 1, 2),
        ("Literature review", 1, 5),
        ("Requirements analysis", 3, 5),
        ("Smart glove circuit design", 5, 7),
        ("Smart stick circuit design", 5, 7),
        ("ESP32 glove firmware", 7, 10),
        ("ESP32 stick firmware", 7, 11),
        ("Android app development", 8, 13),
        ("Camera backup and user signs", 11, 14),
        ("Integration and calibration", 13, 15),
        ("Testing and debugging", 14, 16),
        ("Report writing and presentation", 12, 16),
    ]
    left = 70
    top = 150
    task_w = 430
    week_w = 80
    row_h = 65
    fnt = font(22)
    draw.rectangle((left, top, left + task_w + 16 * week_w, top + row_h), fill=BLUE, outline=LINE, width=2)
    draw.text((left + 20, top + 20), "Activity", fill=INK, font=font(22, bold=True))
    for w in range(1, 17):
        x = left + task_w + (w - 1) * week_w
        draw.text((x + 20, top + 20), str(w), fill=INK, font=font(22, bold=True))
    for i, (task, start, end) in enumerate(tasks, 1):
        y = top + i * row_h
        draw.rectangle((left, y, left + task_w + 16 * week_w, y + row_h), outline="#667085", width=1)
        draw.text((left + 15, y + 18), task, fill=INK, font=fnt)
        for w in range(1, 17):
            x = left + task_w + (w - 1) * week_w
            draw.line((x, y, x, y + row_h), fill="#D0D5DD", width=1)
        x1 = left + task_w + (start - 1) * week_w + 8
        x2 = left + task_w + end * week_w - 8
        draw.rounded_rectangle((x1, y + 15, x2, y + 50), radius=12, fill="#1F7A8C")
    path = FIG_DIR / "figure_13_workplan_gantt.png"
    img.save(path)
    return path


def set_cell_margin(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade(cell, color=BLUE.replace("#", "")):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_para(doc, text="", style=None, bold=False, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_caption(doc, text):
    p = add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_margin(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    return table


def add_captioned_table(doc, caption, headers, rows, widths=None, body_font=9.6, header_font=10.0):
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    caption_cell = table.rows[0].cells[0]
    if len(headers) > 1:
        caption_cell = caption_cell.merge(table.rows[0].cells[-1])
    set_cell_margin(caption_cell)
    shade(caption_cell, color="EAF2F8")
    p = caption_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)

    hdr = table.rows[1]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(header_font)

    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_margin(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.name = "Times New Roman"
            r.font.size = Pt(body_font)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    return table


def add_split_captioned_table(doc, caption, headers, rows, chunk_size, widths=None):
    for start in range(0, len(rows), chunk_size):
        if start:
            doc.add_page_break()
        suffix = "" if start == 0 else " (continued)"
        add_captioned_table(doc, caption.replace(".", f"{suffix}."), headers, rows[start:start + chunk_size], widths=widths)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def apply_text_cleanup(doc):
    replacements = {
        "SignalSense.apk": "SignaSense.apk",
        "NOTE: Insert final questionnaire used for user evaluation.": "The detailed user evaluation questionnaire is provided in Appendix M.",
        "NOTE: Insert interview guide questions if formal interviews were conducted.": "The interview guide questions are provided in Appendix N.",
        "NOTE: Insert final Gantt chart or project work plan if required by the department.": "The completed project work plan and Gantt chart are provided in Appendix I.",
    }
    for p in iter_all_paragraphs(doc):
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def insert_signature_table(doc):
    target = None
    for p in doc.paragraphs:
        if p.text.startswith("We declare that this report is our original work"):
            target = p
            break
    if target is None:
        return
    after = target
    intro = after.insert_paragraph_before("")
    after._p.addnext(intro._p)
    intro.add_run("Student declaration signatures").bold = True
    intro.paragraph_format.line_spacing = 1.5

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Name", "Registration Number", "Signature", "Date"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell)
        cell.paragraphs[0].add_run(h).bold = True
    for name, reg in GROUP_MEMBERS:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = reg
        cells[2].text = "________________________"
        cells[3].text = "________________"
    for row in table.rows:
        for cell in row.cells:
            set_cell_margin(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
    intro._p.addnext(table._tbl)


def insert_body_figures(doc, figures):
    captions = {
        "Figure 1: SignaSense conceptual framework.": figures[1],
        "Figure 2: SignaSense integrated system architecture.": figures[2],
        "Figure 3: Smart glove flex-sensor voltage-divider wiring.": figures[3],
        "Figure 5: Smart stick ultrasonic sensor and buzzer wiring.": figures[5],
        "Figure 7: Smart stick left/right scanning procedure for route guidance.": figures[7],
        "Figure 8: SignaSense Android onboarding, accessibility mode and theme selection.": figures[8],
        "Figure 9: Camera sign backup recognition pipeline.": figures[9],
        "Figure 10: User-defined sign training workflow.": figures[10],
    }
    body_started = False
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text == "CHAPTER ONE: INTRODUCTION":
            body_started = True
        if not body_started:
            continue
        if text in captions:
            break_p = p.insert_paragraph_before()
            break_p.add_run().add_break(WD_BREAK.PAGE)
            pic_p = p.insert_paragraph_before()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(str(captions[text]), width=Inches(5.9))


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    new_para.paragraph_format.line_spacing = 1.5
    return new_para


def append_list_entries(doc):
    replacements = {
        "Figure 9: Smart glove BLE screen with live letter, word and sentence suggestions.": "Figure 9: Camera sign backup recognition pipeline.",
        "Figure 9: Smart glove BLE screen showing live letter, word and sentence suggestions.": "Figure 9: Camera sign backup recognition pipeline.",
        "Figure 10: Camera sign backup and user-defined alphabet training interface.": "Figure 10: User-defined sign training workflow.",
    }
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if stripped in replacements:
            p.text = replacements[stripped]
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                for r in p.runs:
                    r.italic = True
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

    for p in list(doc.paragraphs):
        if p.text.strip() == "Figure 10: User-defined sign training workflow.":
            anchor = p
            for entry in [
                "Figure 11: Smart stick distance filtering and guidance logic.",
                "Figure 12: Prototype verification and testing workflow.",
                "Figure 13: SignaSense project work plan Gantt chart.",
            ]:
                anchor = insert_paragraph_after(anchor, entry)
            break
    for p in list(doc.paragraphs):
        if p.text.strip() == "Table 7: Project limitations and mitigation strategies.":
            anchor = p
            for entry in [
                "Table 8: Detailed system requirements matrix.",
                "Table 9: Prototype component and cost estimate.",
                "Table 10: Project work plan.",
                "Table 11: Risk register.",
                "Table 12: Calibration record sheet.",
                "Table 13: Extended functional test cases.",
                "Table 14: User evaluation questionnaire.",
            ]:
                anchor = insert_paragraph_after(anchor, entry)
            break


def add_detailed_appendices(doc, figures):
    doc.add_page_break()
    add_heading(doc, "Appendix F: Generated System Design Figures", 2)
    for idx, caption in [
        (11, "Figure 11: Smart stick distance filtering and guidance logic."),
        (12, "Figure 12: Prototype verification and testing workflow."),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(figures[idx]), width=Inches(6.3))
        add_caption(doc, caption)

    doc.add_page_break()
    add_heading(doc, "Appendix G: Detailed System Requirements Matrix", 2)
    add_para(doc, "This matrix expands the functional and non-functional requirements used during implementation. It can be used during marking to trace each prototype feature to a user need and to a verification method.")
    requirements = [
        ("REQ-01", "Blind user onboarding", "The app shall ask whether the user is blind or deaf at startup.", "MainActivity onboarding with voice prompt.", "Open app and confirm spoken/visual prompt."),
        ("REQ-02", "Language selection", "The app shall offer English, Luganda and Acholi language choices.", "Language buttons and voice command flow.", "Select each language option and verify status text."),
        ("REQ-03", "Theme selection", "The app shall allow dark and light mode selection.", "ThemeSettings and resource colours.", "Toggle themes and confirm readable interface."),
        ("REQ-04", "Glove sensor reading", "The glove shall read five flex sensors continuously.", "analogRead on configured ESP32 pins.", "View raw values in Serial Monitor and BLE status."),
        ("REQ-05", "Glove calibration", "The glove shall calibrate against the user's open hand.", "Open-hand baseline and dynamic bend scale.", "Recalibrate and observe bend percentage reset."),
        ("REQ-06", "Letter stability", "The glove shall avoid accepting unstable letters.", "Stable-letter and release-to-accept workflow.", "Move fingers quickly and confirm random letters are not committed."),
        ("REQ-07", "Word formation", "Detected letters shall form editable words.", "Current word and app-side suggestions.", "Hold signs, accept letters and view word output."),
        ("REQ-08", "Committed speech", "The phone shall speak only the committed word.", "Commit Word and Speak button/command.", "Build a word and press commit."),
        ("REQ-09", "Camera backup", "The app shall detect hand signs using the phone camera.", "CameraX and MediaPipe hand landmarks.", "Open Camera Signs and hold hand in frame."),
        ("REQ-10", "User-defined signs", "A user shall train a private alphabet on their own phone.", "Local CSV storage in private app files.", "Train samples and switch to user-defined mode."),
        ("REQ-11", "Obstacle distance", "The stick shall measure obstacle distance continuously.", "Ultrasonic trigger/echo and filtered samples.", "Place objects at known distances and check serial output."),
        ("REQ-12", "Bluetooth speaker output", "The stick shall send voice to SINOBAND Book.", "ESP32 A2DP source and embedded speech clips.", "Power speaker near stick and listen for ready phrase."),
        ("REQ-13", "No repeated speech", "The stick shall not repeat the same clear reading continuously.", "Speech throttle, bucket state and distance delta logic.", "Hold a fixed obstacle and confirm speech spacing."),
        ("REQ-14", "Urgent stop", "Very close obstacles shall trigger stop guidance.", "Danger thresholds and speech queue.", "Move obstacle below scan trigger distance."),
        ("REQ-15", "Path suggestion", "The stick shall compare left and right scans.", "Top-sample averaging and side comparison.", "Scan left/right and confirm clearer-side advice."),
        ("REQ-16", "Backup buzzer", "The buzzer shall alert when Bluetooth speaker is not connected.", "Distance-based buzzer patterns.", "Disconnect speaker and move obstacle."),
    ]
    add_split_captioned_table(doc, "Table 8: Detailed system requirements matrix.", ["ID", "Requirement", "Description", "Implementation", "Verification"], requirements, chunk_size=8)

    doc.add_page_break()
    add_heading(doc, "Appendix H: Prototype Component and Cost Estimate", 2)
    add_para(doc, "The project was designed around affordable and locally obtainable parts. The costs below are estimates and should be replaced with actual receipts before final financial submission.")
    cost_rows = [
        ("ESP32 development board", "2", "35,000", "70,000", "One board for glove and one for stick."),
        ("Flex sensor", "5", "25,000", "125,000", "One for each finger."),
        ("Fixed resistors", "5+", "500", "3,000", "Voltage-divider resistors and spares."),
        ("Ultrasonic sensor", "1", "12,000", "12,000", "Obstacle distance sensing."),
        ("Buzzer", "1", "3,000", "3,000", "Backup alert output."),
        ("Bluetooth speaker", "1", "Existing", "0", "SINOBAND Book used during prototype tests."),
        ("Jumper wires", "1 set", "10,000", "10,000", "Prototype wiring."),
        ("Breadboard/perfboard", "1", "15,000", "15,000", "Circuit mounting."),
        ("Glove material", "1", "15,000", "15,000", "Wearable base."),
        ("Walking stick frame", "1", "25,000", "25,000", "Prototype stick body."),
        ("Power bank/battery", "2", "30,000", "60,000", "Portable power."),
        ("Enclosure and fasteners", "1 set", "20,000", "20,000", "Protection and mounting."),
    ]
    add_captioned_table(doc, "Table 9: Prototype component and cost estimate.", ["Item", "Qty", "Unit Cost (UGX)", "Estimated Cost (UGX)", "Purpose"], cost_rows)
    add_para(doc, "Estimated subtotal: UGX 358,000. NOTE: Replace this estimate with actual purchased-item values if receipts are available.")

    doc.add_page_break()
    add_heading(doc, "Appendix I: Project Work Plan", 2)
    add_para(doc, "The work plan is structured as a 16-week final-year project schedule. It covers investigation, design, implementation, integration, testing, documentation and presentation.")
    work_rows = [
        ("1-2", "Problem identification and topic refinement", "Define disability problem, project title and scope.", "Approved project direction."),
        ("1-5", "Literature review", "Study smart gloves, smart sticks, accessibility apps and relevant laws.", "Literature review chapter."),
        ("3-5", "Requirements analysis", "Identify user, hardware, software and safety requirements.", "Requirements matrix."),
        ("5-7", "Hardware design", "Draw glove and stick wiring; select pins and components.", "Circuit plan and parts list."),
        ("7-10", "Glove firmware", "Read flex sensors, calibrate, classify bends and send BLE data.", "BLE glove sketch."),
        ("7-11", "Stick firmware", "Read ultrasonic distance, calibrate, speak guidance and scan sides.", "Smart stick sketch."),
        ("8-13", "Android app", "Build onboarding, BLE screen, camera mode and TTS.", "SignaSense APK."),
        ("11-14", "Camera backup and user-defined signs", "Integrate MediaPipe and local training.", "Camera Signs screen."),
        ("13-15", "Integration", "Connect glove to app; connect stick to speaker; refine thresholds.", "Integrated prototype."),
        ("14-16", "Testing and debugging", "Run functional tests and document issues.", "Test summary."),
        ("12-16", "Report writing", "Prepare final report, figures and presentation material.", "Final report."),
    ]
    add_captioned_table(doc, "Table 10: Project work plan.", ["Week(s)", "Activity", "Description", "Deliverable"], work_rows)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(figures[13]), width=Inches(6.5))
    add_caption(doc, "Figure 13: SignaSense project work plan Gantt chart.")

    doc.add_page_break()
    add_heading(doc, "Appendix J: Risk Register", 2)
    risk_rows = [
        ("R1", "Flex sensors read zero", "High", "High", "Check voltage divider wiring, common ground and correct ADC pins."),
        ("R2", "Wi-Fi blocks ADC2 readings", "High", "High", "Use BLE for glove communication and avoid Wi-Fi while reading ADC2 pins."),
        ("R3", "Random letters appear", "High", "Medium", "Use smoothing, stable hold, open-hand release and app-side suggestion filtering."),
        ("R4", "Ultrasonic echo is 5 V", "High", "Medium", "Use voltage divider or level shifter before ESP32 ECHO input."),
        ("R5", "Bluetooth speaker fails to connect", "Medium", "Medium", "Keep speaker close, unpair from phone/PC and use known name/address fallback."),
        ("R6", "Camera mode misreads signs", "Medium", "Medium", "Use clear lighting, full hand in frame and user-defined training option."),
        ("R7", "Battery power is unstable", "High", "Medium", "Use reliable power bank and secure cable connections."),
        ("R8", "Prototype wiring becomes loose", "Medium", "High", "Move from breadboard to soldered/perfboard assembly before field use."),
        ("R9", "User cannot make standard signs", "Medium", "Medium", "Use local user-defined signs stored on phone."),
        ("R10", "Formal testing data unavailable", "Medium", "High", "Document prototype tests and schedule supervised user evaluation."),
    ]
    add_split_captioned_table(doc, "Table 11: Risk register.", ["ID", "Risk", "Impact", "Likelihood", "Mitigation"], risk_rows, chunk_size=5)

    doc.add_page_break()
    add_heading(doc, "Appendix K: Calibration Guide and Record Sheet", 2)
    add_para(doc, "Calibration is required because flex sensors, resistor values, hand size, power supply and ultrasonic sensor mounting can change readings. The tables below should be filled whenever the prototype is prepared for demonstration or field testing.")
    for module, rows in [
        ("Smart glove calibration steps", [
            ("1", "Connect all flex sensors and confirm common ground.", "All raw values change when fingers move."),
            ("2", "Open the SignaSense app and connect to SignaSenseGlove.", "BLE status shows connected."),
            ("3", "Hold the hand open and press Recalibrate Open Hand.", "Baseline starts from current open hand."),
            ("4", "Bend each finger individually.", "The matching bend percentage changes."),
            ("5", "Hold a known letter until stable.", "The app shows stable/pending letter."),
            ("6", "Open hand to accept letter.", "Letter moves into current word."),
        ]),
        ("Smart stick calibration steps", [
            ("1", "Set air temperature in the firmware if known.", "Speed-of-sound calculation is closer to real conditions."),
            ("2", "Place a flat object at 50 cm.", "Serial distance is near 0.5 m."),
            ("3", "Place a flat object at 100 cm.", "Serial distance is near 1.0 m."),
            ("4", "Adjust calibration scale/offset if readings are consistently high or low.", "Measured output matches ruler within acceptable margin."),
            ("5", "Test warning, danger and very-close thresholds.", "Voice guidance changes by range."),
            ("6", "Run left/right scan test.", "Best side is selected correctly."),
        ]),
    ]:
        module_p = add_para(doc, module, bold=True)
        module_p.paragraph_format.keep_with_next = True
        add_table(doc, ["Step", "Action", "Expected Output"], rows)
    add_captioned_table(doc, "Table 12: Calibration record sheet.", ["Date", "Module", "Measured Point", "Observed Reading", "Adjustment Made", "Tester"], [("____", "____", "____", "____", "____", "____") for _ in range(8)])

    doc.add_page_break()
    add_heading(doc, "Appendix L: Extended Functional Test Cases", 2)
    test_cases = []
    for i, (area, case, expected) in enumerate([
        ("Glove", "Raw sensor values print clearly to Serial Monitor.", "Five raw values are visible or no-signal state is reported."),
        ("Glove", "Open-hand recalibration resets baseline.", "Bend percentages return near straight when hand is open."),
        ("Glove", "Stable letter hold creates pending letter.", "Pending letter appears only after stable hold."),
        ("Glove", "Open hand accepts pending letter.", "Letter is appended to word."),
        ("Glove", "Commit Word and Speak outputs one full word.", "Phone speaks the committed word once."),
        ("Glove", "Clear Word resets current word.", "Word field becomes empty."),
        ("App", "Blind onboarding speaks first prompt.", "Phone says SignaSense is ready."),
        ("App", "Deaf mode shows visual controls.", "Visual dashboard appears with glove/camera options."),
        ("App", "Dark/light mode persists.", "Theme remains after activity changes."),
        ("Camera", "General signs detect A-Z shape.", "Letter appears after stable hold."),
        ("Camera", "No hand clears raw letter.", "Letter returns to idle after release."),
        ("Camera", "User-defined training captures samples.", "Sample count increases to eight per letter."),
        ("Stick", "Startup buzzer test runs.", "Two short beeps occur."),
        ("Stick", "Bluetooth speaker connects.", "Smart stick ready phrase plays."),
        ("Stick", "Obstacle at warning distance triggers guidance.", "Distance and slow-down/obstacle phrase plays."),
        ("Stick", "Very close obstacle triggers stop.", "Obstacle ahead and stop now are spoken."),
        ("Stick", "Left/right scan compares clear sides.", "Clearer side is recommended."),
        ("Stick", "Speaker disconnected uses buzzer.", "Distance-pattern buzzer output works."),
        ("Stick", "Invalid echo is rejected.", "No false distance is spoken."),
        ("Integration", "Phone app and glove communicate while stick runs independently.", "Both assistive modules operate without blocking each other."),
    ], 1):
        test_cases.append((f"TC-{i:02d}", area, case, expected, "Pass/Fail: ____", "Notes: __________________"))
    add_split_captioned_table(doc, "Table 13: Extended functional test cases.", ["ID", "Area", "Test Case", "Expected Result", "Result", "Notes"], test_cases, chunk_size=8)

    doc.add_page_break()
    add_heading(doc, "Appendix M: User Evaluation Questionnaire", 2)
    add_para(doc, "This questionnaire should be completed after supervised prototype testing. Use a five-point scale: 1 = strongly disagree, 2 = disagree, 3 = neutral, 4 = agree, 5 = strongly agree.")
    questionnaire = [
        ("Q1", "The smart glove was comfortable to wear."),
        ("Q2", "The glove detected deliberate signs more accurately after calibration."),
        ("Q3", "The word suggestion section helped reduce typing or correction effort."),
        ("Q4", "The phone spoke committed words clearly."),
        ("Q5", "The camera backup was easy to understand."),
        ("Q6", "The user-defined training option would help users with different hand shapes."),
        ("Q7", "The smart stick gave obstacle warnings early enough."),
        ("Q8", "The spoken distance was understandable."),
        ("Q9", "The left/right scan guidance was easy to follow."),
        ("Q10", "The backup buzzer was useful when the Bluetooth speaker was not connected."),
        ("Q11", "The app interface was readable in light mode."),
        ("Q12", "The app interface was readable in dark mode."),
        ("Q13", "The system could improve communication independence."),
        ("Q14", "The system could improve mobility confidence."),
        ("Q15", "I would recommend improving this prototype for real-world use."),
    ]
    add_split_captioned_table(doc, "Table 14: User evaluation questionnaire.", ["No.", "Statement", "1", "2", "3", "4", "5"], [(q, s, "", "", "", "", "") for q, s in questionnaire], chunk_size=5)

    doc.add_page_break()
    add_heading(doc, "Appendix N: Interview Guide", 2)
    interview_sections = [
        ("Background", [
            "What assistive devices do you currently use for communication or mobility?",
            "What challenges do you face when using those devices?",
            "How often do you need help from another person during communication or movement?",
        ]),
        ("Smart glove feedback", [
            "Were the glove sensors comfortable and secure?",
            "Which signs were easy or difficult to form?",
            "Would word suggestions help prevent wrong speech output?",
            "What local-language or common phrase features should be added?",
        ]),
        ("Smart stick feedback", [
            "Was the spoken distance clear and useful?",
            "Was the stop-now warning early enough?",
            "Was the left/right scanning instruction easy to follow?",
            "Would vibration feedback also be useful?",
        ]),
        ("App feedback", [
            "Was blind/deaf onboarding clear?",
            "Was dark/light mode useful?",
            "Would you prefer voice commands, buttons or both?",
            "How should the app improve captions and suggested sentences?",
        ]),
    ]
    for title, qs in interview_sections:
        add_para(doc, title, bold=True)
        for q in qs:
            add_para(doc, f"- {q}")

    doc.add_page_break()
    add_heading(doc, "Appendix O: Prototype User Manual", 2)
    manual_sections = [
        ("Using the smart glove", [
            "Charge or power the ESP32 glove before use.",
            "Open the SignaSense app and choose the deaf/visual interface or smart glove option.",
            "Tap Connect Smart Glove BLE and wait for SignaSenseGlove.",
            "Hold the hand open and press Recalibrate Open Hand.",
            "Make one letter sign and hold it steady until the app marks it stable or pending.",
            "Open the hand briefly to accept the letter into the current word.",
            "Use suggestions if the intended word appears.",
            "Press Commit Word and Speak only after the full word is correct.",
        ]),
        ("Using camera sign backup", [
            "Open Camera Signs in the app.",
            "Place the full hand in the camera frame with good lighting.",
            "Use General Signs for the built-in alphabet or User-Defined Signs for trained signs.",
            "For user-defined signs, capture enough samples for each letter A to Z.",
            "Commit completed words before speaking.",
        ]),
        ("Using the smart stick", [
            "Turn on the Bluetooth speaker named SINOBAND Book and keep it close to the stick.",
            "Power the ESP32 smart stick.",
            "Wait for the ready phrase.",
            "Move forward only when the stick is not warning of an obstacle.",
            "When it says stop now, stop first and then follow scan instructions.",
            "Scan left slowly when instructed, then scan right slowly.",
            "Move only toward the side recommended by the stick if the environment is safe.",
        ]),
        ("Safety notes", [
            "The prototype is not a certified medical or mobility device.",
            "Use it with supervision during demonstrations and testing.",
            "Do not rely on it near roads, stairs, open drains, water or hazardous machinery.",
            "Protect ESP32 pins from 5 V sensor signals using voltage division or level shifting.",
        ]),
    ]
    for title, items in manual_sections:
        add_para(doc, title, bold=True)
        for item in items:
            add_para(doc, f"- {item}")

    doc.add_page_break()
    add_heading(doc, "Appendix P: Maintenance and Future Improvement Notes", 2)
    notes = [
        "Move prototype wiring from breadboard and jumper wires to soldered perfboard or a printed circuit board before field use.",
        "Add strain relief to flex sensor leads to prevent wire breakage at the fingers.",
        "Use a stable 3.3 V supply for the glove sensors and avoid loose USB or power-bank cables.",
        "Create a protective enclosure for the smart stick ESP32 and ultrasonic sensor.",
        "Add more obstacle sensors or a servo-mounted sensor to reduce the need for manual left/right sweeping.",
        "Record high-quality local language voice clips for English, Luganda and Acholi prompts.",
        "Expand the glove beyond alphabet letters by adding phrase gestures and better motion sensing.",
        "Improve the camera model using real training data collected from the target users.",
        "Conduct field trials with users and update thresholds based on measured accuracy and comfort.",
    ]
    for note in notes:
        add_para(doc, f"- {note}")


def main():
    figures = generate_figures()
    doc = Document(BASE_DOCX)
    insert_signature_table(doc)
    apply_text_cleanup(doc)
    append_list_entries(doc)
    insert_body_figures(doc, figures)
    add_detailed_appendices(doc, figures)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
